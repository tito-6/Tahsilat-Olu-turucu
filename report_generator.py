"""
Report generator module for Excel, PDF, and Word output
Generates various payment reports with formatting
"""

import pandas as pd
import xlsxwriter
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path

from data_import import PaymentData
from currency import convert_payment_to_usd
# Removed currency_optimizer imports - using PaymentData.usd_amount directly

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates various payment reports in multiple formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
        self.optimized_payments = None
    
    def setup_custom_styles(self):
        """Setup custom styles for reports with Turkish character support"""
        # Register Turkish-friendly fonts
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        try:
            # Try to register Arial font for better Turkish support
            pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
            pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
            font_name = 'Arial'
            font_bold = 'Arial-Bold'
        except:
            # Fallback to DejaVu fonts if Arial not available
            try:
                pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
                font_name = 'DejaVuSans'
                font_bold = 'DejaVuSans-Bold'
            except:
                # Final fallback to default fonts
                font_name = 'Helvetica'
                font_bold = 'Helvetica-Bold'
        
        # Custom paragraph styles with Turkish character support
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            fontName=font_bold,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        self.subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=self.styles['Heading2'],
            fontSize=12,
            fontName=font_bold,
            spaceAfter=12,
            alignment=1  # Center
        )
        
        self.header_style = ParagraphStyle(
            'CustomHeader',
            parent=self.styles['Normal'],
            fontSize=10,
            fontName=font_bold,
            alignment=1  # Center
        )
        
        self.normal_style = ParagraphStyle(
            'CustomNormal',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName=font_name
        )
    
    def _ensure_turkish_encoding(self, text):
        """Ensure text is properly encoded for Turkish characters"""
        if isinstance(text, str):
            # Ensure proper UTF-8 encoding
            return text.encode('utf-8').decode('utf-8')
        return str(text) if text is not None else ""
    
    def _create_turkish_test_data(self):
        """Create test data with Turkish characters for verification"""
        return {
            'customer': 'Müşteri Adı Soyadı',
            'project': 'Proje Adı',
            'account': 'Hesap Adı',
            'payment_type': 'BANK_TRANSFER',
            'currency': 'TL',
            'description': 'Açıklama ve detaylar'
        }
    
    def generate_daily_usd_breakdown(self, payments: List[PaymentData], 
                                   start_date: datetime, end_date: datetime) -> pd.DataFrame:
        """Generate daily USD breakdown by client and project"""
        # Filter payments by date range
        filtered_payments = [p for p in payments 
                           if p.date and start_date <= p.date <= end_date]
        
        # Use already converted USD amounts
        processed_payments = []
        for payment in filtered_payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            processed_payments.append({
                'customer': payment.customer_name,
                'project': payment.project_name,
                'date': payment.date,
                'amount_usd': usd_amount,
                'original_amount': payment.original_amount,
                'currency': payment.currency,
                'exchange_rate': payment.conversion_rate if payment.conversion_rate else 1.0
            })
        
        if not processed_payments:
            return pd.DataFrame()
        
        # Create pivot table
        df = pd.DataFrame(processed_payments)
        df['date_str'] = df['date'].dt.strftime('%Y-%m-%d')
        df['day_name'] = df['date'].dt.day_name()
        
        # Pivot by customer, project, and date
        pivot = df.pivot_table(
            index=['customer', 'project'],
            columns='date_str',
            values='amount_usd',
            aggfunc='sum',
            fill_value=0
        )
        
        # Add total column
        pivot['Genel Toplam'] = pivot.sum(axis=1)
        
        return pivot
    
    def generate_weekly_summary(self, payments: List[PaymentData]) -> pd.DataFrame:
        """Generate weekly summary by project"""
        if not payments:
            return pd.DataFrame()
        
        # Process payments
        processed_payments = []
        for payment in payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            if payment.date:
                # Get week start (Monday)
                week_start = payment.date - timedelta(days=payment.date.weekday())
                processed_payments.append({
                    'project': payment.project_name,
                    'week_start': week_start,
                    'amount_usd': usd_amount
                })
        
        if not processed_payments:
            return pd.DataFrame()
        
        df = pd.DataFrame(processed_payments)
        df['week_str'] = df['week_start'].dt.strftime('%Y-W%U')
        
        # Group by project and week
        summary = df.groupby(['project', 'week_str'])['amount_usd'].sum().unstack(fill_value=0)
        
        return summary
    
    def generate_monthly_summary(self, payments: List[PaymentData]) -> pd.DataFrame:
        """Generate monthly summary by project and payment type (BANK_TRANSFER, Nakit, Çek)"""
        if not payments:
            return pd.DataFrame()
        
        # Process payments with proper payment type classification
        processed_payments = []
        for payment in payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            if payment.date:
                month_str = payment.date.strftime('%Y-%m')
                
                # Determine payment type using the same logic as weekly tables
                payment_type = self._classify_payment_type(payment)
                
                processed_payments.append({
                    'project': payment.project_name,
                    'payment_type': payment_type,
                    'month': month_str,
                    'amount_usd': usd_amount
                })
        
        if not processed_payments:
            return pd.DataFrame()
        
        df = pd.DataFrame(processed_payments)
        
        # Create pivot table with payment types as rows and projects as columns
        pivot = df.pivot_table(
            index='payment_type',
            columns='project',
            values='amount_usd',
            aggfunc='sum',
            fill_value=0
        )
        
        # Add total row
        pivot.loc['TOPLAM'] = pivot.sum()
        
        return pivot
    
    def _classify_payment_type(self, payment: PaymentData) -> str:
        """Classify payment type using the same logic as weekly tables"""
        # Check for check payments first
        if payment.is_check_payment:
            return 'Çek'
        
        # Check tahsilat_sekli field
        if payment.tahsilat_sekli:
            tahsilat_upper = payment.tahsilat_sekli.upper()
            if 'NAKİT' in tahsilat_upper or 'NAKIT' in tahsilat_upper:
                return 'Nakit'
            elif 'BANKA' in tahsilat_upper or 'HAVALE' in tahsilat_upper:
                return 'BANK_TRANSFER'
            elif 'ÇEK' in tahsilat_upper or 'CEK' in tahsilat_upper:
                return 'Çek'
        
        # Check account name as fallback
        account_upper = payment.account_name.upper() if payment.account_name else ''
        
        # Check for Yapı Kredi and other bank accounts
        if any(keyword in account_upper for keyword in [
            'YAPI KREDİ', 'YAPI KREDI', 'YAPIKREDÝ', 'YAPIKREDI', 'YAPI',
            'HAVALE', 'TRANSFER', 'BANKA', 'GARANTI', 'İŞ BANKASI'
        ]):
            return 'BANK_TRANSFER'
        elif 'KASA' in account_upper and 'NAKİT' not in account_upper:
            return 'Nakit'
        elif any(keyword in account_upper for keyword in [
            'NAKİT', 'NAKIT', 'NAKÝT', 'CASH'
        ]):
            return 'Nakit'
        
        # Default to BANK_TRANSFER for TL payments from bank accounts
        if payment.is_tl_payment and account_upper:
            return 'BANK_TRANSFER'
        
        # Default fallback
        return 'BANK_TRANSFER'
    
    def generate_daily_timeline(self, payments: List[PaymentData], 
                              start_date: datetime, end_date: datetime) -> pd.DataFrame:
        """Generate daily USD timeline across the month"""
        # Filter payments by date range
        filtered_payments = [p for p in payments 
                           if p.date and start_date <= p.date <= end_date]
        
        # Process payments
        processed_payments = []
        for payment in filtered_payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            processed_payments.append({
                'date': payment.date,
                'amount_usd': usd_amount,
                'project': payment.project_name
            })
        
        if not processed_payments:
            return pd.DataFrame()
        
        df = pd.DataFrame(processed_payments)
        df['date_str'] = df['date'].dt.strftime('%Y-%m-%d')
        
        # Group by date and project
        timeline = df.groupby(['date_str', 'project'])['amount_usd'].sum().unstack(fill_value=0)
        
        # Add total column
        timeline['Günlük Toplam'] = timeline.sum(axis=1)
        
        return timeline
    
    def generate_payment_type_summary(self, payments: List[PaymentData]) -> pd.DataFrame:
        """Generate TL and USD totals by payment type"""
        if not payments:
            return pd.DataFrame()
        
        # Process payments
        processed_payments = []
        for payment in payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            processed_payments.append({
                'channel': payment.payment_channel,
                'amount_tl': payment.original_amount if payment.is_tl_payment else 0,
                'amount_usd': usd_amount,
                'currency': payment.currency
            })
        
        if not processed_payments:
            return pd.DataFrame()
        
        df = pd.DataFrame(processed_payments)
        
        # Group by channel
        summary = df.groupby('channel').agg({
            'amount_tl': 'sum',
            'amount_usd': 'sum'
        }).round(2)
        
        # Add total row
        summary.loc['TOPLAM'] = summary.sum()
        
        return summary
    
    def generate_customer_date_table(self, payments: List[PaymentData], 
                                   start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Generate customer-date pivot table with weekly separation and TL payment tracking"""
        # Filter payments by date range
        filtered_payments = [p for p in payments 
                           if p.date and start_date <= p.date <= end_date]
        
        if not filtered_payments:
            return {}
        
        # Convert TL payments to USD and track conversion info with rates
        processed_payments = []
        for payment in filtered_payments:
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            is_converted = payment.is_tl_payment and payment.usd_amount > 0
            conversion_rate = payment.conversion_rate if payment.conversion_rate else 1.0
            
            processed_payments.append({
                'customer': payment.customer_name,
                'project': payment.project_name,
                'date': payment.date,
                'amount_usd': usd_amount,
                'original_amount': payment.original_amount,
                'currency': payment.currency,
                'is_tl_converted': is_converted,
                'conversion_rate': conversion_rate,
                'week_start': payment.date - timedelta(days=payment.date.weekday())  # Monday of the week
            })
        
        if not processed_payments:
            return {}
        
        # Group by weeks
        df = pd.DataFrame(processed_payments)
        weeks = df.groupby('week_start')
        
        weekly_tables = {}
        
        # Turkish day names mapping
        turkish_days = {
            'Monday': 'Pazartesi',
            'Tuesday': 'Salı', 
            'Wednesday': 'Çarşamba',
            'Thursday': 'Perşembe',
            'Friday': 'Cuma',
            'Saturday': 'Cumartesi',
            'Sunday': 'Pazar'
        }
        
        for week_start, week_data in weeks:
            # Create complete week structure (Monday to Sunday)
            week_dates = []
            week_day_names = []
            
            for day_offset in range(7):  # Monday (0) to Sunday (6)
                current_date = week_start + timedelta(days=day_offset)
                if start_date <= current_date <= end_date:
                    date_str = current_date.strftime('%d.%m.%Y')
                    day_name_english = current_date.strftime('%A')
                    day_name_turkish = turkish_days.get(day_name_english, day_name_english)
                    
                    week_dates.append(date_str)
                    # Use actual date with weekday name for better clarity
                    week_day_names.append(f"{day_name_turkish}<br>{date_str}")
            
            # Format date as DD.MM.YYYY for columns
            week_data['date_str'] = week_data['date'].dt.strftime('%d.%m.%Y')
            
            # Create pivot table for this week with all 7 days
            pivot = week_data.pivot_table(
                index=['customer', 'project'],
                columns='date_str',
                values='amount_usd',
                aggfunc='sum',
                fill_value=0
            )
            
            # Create TL conversion tracking table
            tl_pivot = week_data.pivot_table(
                index=['customer', 'project'],
                columns='date_str',
                values='is_tl_converted',
                aggfunc='any',
                fill_value=False
            )
            
            # Create conversion rate tracking table
            rate_pivot = week_data.pivot_table(
                index=['customer', 'project'],
                columns='date_str',
                values='conversion_rate',
                aggfunc='first',
                fill_value=0
            )
            
            # Ensure all week dates are present as columns (fill missing days with 0)
            for date_str in week_dates:
                if date_str not in pivot.columns:
                    pivot[date_str] = 0
                if date_str not in tl_pivot.columns:
                    tl_pivot[date_str] = False
                if date_str not in rate_pivot.columns:
                    rate_pivot[date_str] = 0
            
            # Sort columns by date and reorder to match week structure
            pivot = pivot[week_dates]
            tl_pivot = tl_pivot[week_dates]
            rate_pivot = rate_pivot[week_dates]
            
            # Sort customers alphabetically
            if not pivot.empty:
                pivot = pivot.sort_index(level=0)  # Sort by customer name
                tl_pivot = tl_pivot.sort_index(level=0)
                rate_pivot = rate_pivot.sort_index(level=0)
                
                # Add total column
                pivot['Genel Toplam'] = pivot.sum(axis=1)
                tl_pivot['Genel Toplam'] = False
                rate_pivot['Genel Toplam'] = 0
            
            # Create day names mapping for this week
            day_names_dict = dict(zip(week_dates, week_day_names))
            
            # Store week data
            week_end = week_start + timedelta(days=6)
            weekly_tables[week_start] = {
                'pivot': pivot,
                'tl_converted': tl_pivot,
                'conversion_rates': rate_pivot,
                'day_names': day_names_dict,
                'week_range': f"{week_start.strftime('%d.%m.%Y')} - {week_end.strftime('%d.%m.%Y')}",
                'week_dates': week_dates,
                'week_day_names': week_day_names
            }
        
        return weekly_tables
    
    def generate_customer_check_table(self, payments: List[PaymentData], 
                                    start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Generate customer-check pivot table with weekly separation and maturity date handling"""
        # Filter payments by payment date range and check payments only
        # Check payments are identified by "Tahsilat Şekli" = "Çek" or having check amount > 0
        # For filtering, we use payment date consistently; maturity date is used only for currency conversion
        filtered_payments = []
        for p in payments:
            if p.date and p.is_check_payment:
                # Use payment date for filtering (consistent with other payment types)
                if start_date <= p.date <= end_date:
                    filtered_payments.append(p)
        
        if not filtered_payments:
            return {}
        
        # Process check payments with maturity date handling
        processed_checks = []
        for payment in filtered_payments:
            # Use check amount, default to payment amount if check amount is 0
            check_amount = payment.cek_tutari if payment.cek_tutari > 0 else payment.amount
            
            # Ensure we have a maturity date
            maturity_date = payment.cek_vade_tarihi
            if not maturity_date:
                # Default to 6 months after payment date
                maturity_date = payment.date + timedelta(days=180)
            
            processed_checks.append({
                'customer': payment.customer_name,
                'project': payment.project_name,
                'date': payment.date,
                'check_amount_tl': check_amount,
                'maturity_date': maturity_date,
                'week_start': payment.date - timedelta(days=payment.date.weekday())  # Monday of the payment week
            })
        
        if not processed_checks:
            return {}
        
        # Group by weeks
        df = pd.DataFrame(processed_checks)
        weeks = df.groupby('week_start')
        
        weekly_check_tables = {}
        
        # Turkish day names mapping
        turkish_days = {
            'Monday': 'Pazartesi',
            'Tuesday': 'Salı', 
            'Wednesday': 'Çarşamba',
            'Thursday': 'Perşembe',
            'Friday': 'Cuma',
            'Saturday': 'Cumartesi',
            'Sunday': 'Pazar'
        }
        
        for week_start, week_data in weeks:
            # Create complete week structure (Monday to Sunday)
            week_dates = []
            week_day_names = []
            
            for day_offset in range(7):  # Monday (0) to Sunday (6)
                current_date = week_start + timedelta(days=day_offset)
                if start_date <= current_date <= end_date:
                    date_str = current_date.strftime('%d.%m.%Y')
                    day_name_english = current_date.strftime('%A')
                    day_name_turkish = turkish_days.get(day_name_english, day_name_english)
                    
                    week_dates.append(date_str)
                    # Use actual date with weekday name for better clarity
                    week_day_names.append(f"{day_name_turkish}<br>{date_str}")
            
            # Format date as DD.MM.YYYY for columns
            week_data['date_str'] = week_data['date'].dt.strftime('%d.%m.%Y')
            
            # Create pivot table for check amounts (TL) - same structure as regular payments
            check_pivot_tl = week_data.pivot_table(
                index=['customer', 'project'],
                columns='date_str',
                values='check_amount_tl',
                aggfunc='sum',
                fill_value=0
            )
            
            # Create USD equivalent table (using maturity date exchange rate)
            week_data['check_amount_usd'] = week_data.apply(
                lambda row: self.convert_tl_to_usd_at_maturity(row['check_amount_tl'], row['maturity_date']), 
                axis=1
            )
            
            check_pivot_usd = week_data.pivot_table(
                index=['customer', 'project'],
                columns='date_str',
                values='check_amount_usd',
                aggfunc='sum',
                fill_value=0
            )
            
            # Ensure all week dates are present as columns
            for date_str in week_dates:
                if date_str not in check_pivot_tl.columns:
                    check_pivot_tl[date_str] = 0
                if date_str not in check_pivot_usd.columns:
                    check_pivot_usd[date_str] = 0
            
            # Sort columns by date and reorder to match week structure
            check_pivot_tl = check_pivot_tl[week_dates]
            check_pivot_usd = check_pivot_usd[week_dates]
            
            # Sort customers alphabetically
            if not check_pivot_tl.empty:
                check_pivot_tl = check_pivot_tl.sort_index(level=0)
                check_pivot_usd = check_pivot_usd.sort_index(level=0)
                
                # Add total column
                check_pivot_tl['Genel Toplam'] = check_pivot_tl.sum(axis=1)
                check_pivot_usd['Genel Toplam'] = check_pivot_usd.sum(axis=1)
            
            # Create day names mapping for this week
            day_names_dict = dict(zip(week_dates, week_day_names))
            
            # Store week data
            week_end = week_start + timedelta(days=6)
            weekly_check_tables[week_start] = {
                'pivot_tl': check_pivot_tl,
                'pivot_usd': check_pivot_usd,
                'day_names': day_names_dict,
                'week_range': f"{week_start.strftime('%d.%m.%Y')} - {week_end.strftime('%d.%m.%Y')}",
                'week_dates': week_dates,
                'week_day_names': week_day_names
            }
        
        return weekly_check_tables
    
    def convert_tl_to_usd_at_maturity(self, tl_amount: float, maturity_date: datetime) -> float:
        """Convert TL amount to USD using exchange rate at maturity date"""
        try:
            from currency import CurrencyConverter
            converter = CurrencyConverter()
            
            # Get exchange rate for maturity date
            rate = converter.get_usd_rate(maturity_date)
            if rate:
                return tl_amount / rate
            else:
                # Fallback to current rate or estimate
                return tl_amount / 30.0  # Default estimate
        except Exception as e:
            logger.warning(f"Failed to get exchange rate for maturity date {maturity_date}: {e}")
            return tl_amount / 30.0  # Default estimate
    
    def generate_html_preview(self, payments: List[PaymentData], 
                            start_date: datetime, end_date: datetime) -> Dict[str, str]:
        """Generate HTML preview for all report sheets"""
        try:
            # Generate reports using already converted USD amounts from PaymentData
            logger.info("Generating HTML preview...")
            
            # Generate reports
            customer_date_table = self.generate_customer_date_table(payments, start_date, end_date)
            customer_check_table = self.generate_customer_check_table(payments, start_date, end_date)
            
            html_sheets = {}
            
            # Generate HTML for each weekly sheet
            if customer_date_table:
                sorted_weeks = sorted(customer_date_table.keys())
                
                for week_num, week_start in enumerate(sorted_weeks, 1):
                    week_data = customer_date_table[week_start]
                    check_data = customer_check_table.get(week_start) if customer_check_table else None
                    
                    html_content = self._generate_week_html(
                        week_num, week_data, check_data, start_date, end_date, payments
                    )
                    
                    # Create date range name instead of generic "Hafta X"
                    week_range = week_data.get('week_range', '')
                    if week_range and '-' in week_range:
                        # Extract start and end dates from week_range
                        try:
                            start_str, end_str = week_range.split(' - ')
                            sheet_name = f"{start_str}-{end_str}"
                        except:
                            sheet_name = f"Hafta {week_num}"
                    else:
                        sheet_name = f"Hafta {week_num}"
                    
                    html_sheets[sheet_name] = html_content
            
            # Add summary sheet
            html_sheets["Özet"] = self._generate_summary_html(payments, start_date, end_date)
            
            return html_sheets
            
        except Exception as e:
            logger.error(f"Failed to generate HTML preview: {e}")
            return {"Hata": f"<p>Rapor önizleme oluşturulamadı: {e}</p>"}
    
    def _generate_week_html(self, week_num: int, week_data: Dict, check_data: Dict, 
                          start_date: datetime, end_date: datetime, payments: List[PaymentData] = None) -> str:
        """Generate HTML for a single week"""
        pivot_table = week_data['pivot']
        week_range = week_data['week_range']
        week_dates = week_data['week_dates']
        week_day_names = week_data['week_day_names']
        
        # Generate analysis data for the full date range (not just this week)
        week_start = week_data.get('week_start')
        if payments:
            # Use the full date range for analysis, not just this week
            # This ensures all payment types/projects/locations are included
            payment_type_analysis = self.generate_payment_type_analysis(payments, start_date, end_date)
            project_totals_analysis = self.generate_project_totals_analysis(payments, start_date, end_date)
            location_analysis = self.generate_location_analysis(payments, start_date, end_date)
        else:
            payment_type_analysis = {'weekly': {}, 'monthly': {}}
            project_totals_analysis = {'weekly': {}, 'monthly': {}}
            location_analysis = {'weekly': {}, 'monthly': {}}
        
        # Build HTML
        html = f"""
        <div style="font-family: Arial, sans-serif; padding: 20px;">
            <h1 style="text-align: center; color: #2c3e50; margin-bottom: 15px; font-size: 28px; font-weight: 800;">
                PAYMENT REPORTING SYSTEM - DATE RANGES TABLE
            </h1>
            <h2 style="text-align: center; color: #34495e; margin-bottom: 20px; font-size: 20px; font-weight: 600;">
                {week_range} | Hafta {week_num}
            </h2>
            <p style="text-align: center; color: #27ae60; margin-bottom: 30px; font-size: 14px; font-weight: 600; 
               background-color: #eafaf1; padding: 8px; border-radius: 6px; border: 1px solid #27ae60;">
                ✓ Raporlarda 'Ödenen Tutar' değerleri ve 'Ödenen Döviz' kurları kullanılmaktadır
            </p>
            
            <h3 style="color: #2980b9; border-bottom: 2px solid #2980b9; padding-bottom: 5px;">
                HAFTANIN TÜM ÖDEMELERİ (Ödenen Tutar Bazında)
            </h3>
            
            <table style="width: 100%; border-collapse: collapse; margin-bottom: 30px; font-size: 12px;">
                <thead>
                    <tr style="background-color: #ecf0f1;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">SIRA NO</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">MÜŞTERİ ADI SOYADI</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJE</th>
        """
        
        # Add day names row
        for day_name in week_day_names:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center; background-color: #d5dbdb;">{day_name}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">GENEL TOPLAM</th></tr>'
        
        # Add date headers
        html += '<tr style="background-color: #ecf0f1;">'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        
        for date_str in week_dates:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{date_str}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th></tr></thead><tbody>'
        
        # Add data rows
        if not pivot_table.empty:
            sira_no = 1
            for (customer, project), row_data in pivot_table.iterrows():
                html += f'<tr style="background-color: {"#f8f9fa" if sira_no % 2 == 0 else "white"};">'
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{sira_no}</td>'
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{customer}</td>'
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{project}</td>'
                
                for date_str in week_dates:
                    amount = row_data.get(date_str, 0) if hasattr(row_data, 'get') else 0
                    if amount > 0:
                        html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>'
                    else:
                        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
                
                total = row_data.get('Genel Toplam', 0) if hasattr(row_data, 'get') else 0
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right; font-weight: bold;">${total:,.2f}</td>'
                html += '</tr>'
                sira_no += 1
            
            # Add total row
            html += '<tr style="background-color: #2980b9; color: white; font-weight: bold;">'
            html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
            html += '<td style="border: 1px solid #bdc3c7; padding: 8px;">HAFTA TOPLAMI</td>'
            html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
            
            for date_str in week_dates:
                total = pivot_table[date_str].sum() if date_str in pivot_table.columns else 0
                if total > 0:
                    html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total:,.2f}</td>'
                else:
                    html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
            
            grand_total = pivot_table['Genel Toplam'].sum()
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${grand_total:,.2f}</td>'
            html += '</tr>'
        
        html += '</tbody></table>'
        
        # Add check payments table if exists
        if check_data and not check_data['pivot_tl'].empty:
            html += self._generate_check_table_html(check_data, week_dates, week_day_names)
        else:
            # Add empty check payments table
            html += self._generate_empty_check_table_html(week_dates, week_day_names)
        
        # Add analysis tables
        html += self._generate_simple_analysis_tables_html(payment_type_analysis, project_totals_analysis, location_analysis, week_start, is_weekly_sheet=True)
        
        html += '</div>'
        return html
    
    def _generate_check_table_html(self, check_data: Dict, week_dates: List[str], week_day_names: List[str]) -> str:
        """Generate HTML for check payments table"""
        check_pivot_tl = check_data['pivot_tl']
        check_pivot_usd = check_data['pivot_usd']
        
        html = f"""
            <h3 style="color: #e67e22; border-bottom: 2px solid #e67e22; padding-bottom: 5px; margin-top: 40px;">
                ÇEK TAHSİLATLARI
            </h3>
            
            <table style="width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 12px;">
                <thead>
                    <tr style="background-color: #fdf2e9;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">SIRA NO</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">MÜŞTERİ ADI SOYADI</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJE</th>
        """
        
        # Add day names
        for day_name in week_day_names:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center; background-color: #f39c12; color: white;">{day_name}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">GENEL TOPLAM</th></tr>'
        
        # Add date headers
        html += '<tr style="background-color: #fdf2e9;">'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        
        for date_str in week_dates:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{date_str}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th></tr></thead><tbody>'
        
        # TL amounts
        sira_no = 1
        for (customer, project), row_data in check_pivot_tl.iterrows():
            html += f'<tr style="background-color: {"#fef9e7" if sira_no % 2 == 0 else "#fff8e1"};">'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{sira_no}</td>'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{customer}</td>'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{project}</td>'
            
            for date_str in week_dates:
                amount = row_data.get(date_str, 0) if hasattr(row_data, 'get') else 0
                if amount > 0:
                    html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{amount:,.2f}</td>'
                else:
                    html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
            
            total = row_data.get('Genel Toplam', 0) if hasattr(row_data, 'get') else 0
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right; font-weight: bold;">₺{total:,.2f}</td>'
            html += '</tr>'
            sira_no += 1
        
        # TL total row
        html += '<tr style="background-color: #f39c12; color: white; font-weight: bold;">'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM TL</td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        for date_str in week_dates:
            total = check_pivot_tl[date_str].sum() if date_str in check_pivot_tl.columns else 0
            if total > 0:
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{total:,.2f}</td>'
            else:
                html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        tl_grand_total = check_pivot_tl['Genel Toplam'].sum()
        html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{tl_grand_total:,.2f}</td>'
        html += '</tr>'
        
        # USD amounts
        sira_no = 1
        for (customer, project), row_data in check_pivot_usd.iterrows():
            html += f'<tr style="background-color: {"#e8f5e8" if sira_no % 2 == 0 else "#f0f8f0"};">'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{sira_no}</td>'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{customer}</td>'
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px;">{project}</td>'
            
            for date_str in week_dates:
                amount = row_data.get(date_str, 0) if hasattr(row_data, 'get') else 0
                if amount > 0:
                    html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>'
                else:
                    html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
            
            total = row_data.get('Genel Toplam', 0) if hasattr(row_data, 'get') else 0
            html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right; font-weight: bold;">${total:,.2f}</td>'
            html += '</tr>'
            sira_no += 1
        
        # USD total row
        html += '<tr style="background-color: #27ae60; color: white; font-weight: bold;">'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM USD (Vade Tarihi Kuru)</td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        for date_str in week_dates:
            total = check_pivot_usd[date_str].sum() if date_str in check_pivot_usd.columns else 0
            if total > 0:
                html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total:,.2f}</td>'
            else:
                html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        usd_grand_total = check_pivot_usd['Genel Toplam'].sum()
        html += f'<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${usd_grand_total:,.2f}</td>'
        html += '</tr>'
        
        html += '</tbody></table>'
        return html
    
    def _generate_empty_check_table_html(self, week_dates: List[str], week_day_names: List[str]) -> str:
        """Generate HTML for empty check payments table"""
        html = f"""
            <h3 style="color: #e67e22; border-bottom: 2px solid #e67e22; padding-bottom: 5px; margin-top: 40px;">
                ÇEK TAHSİLATLARI
            </h3>
            
            <table style="width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 12px;">
                <thead>
                    <tr style="background-color: #fdf2e9;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">SIRA NO</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">MÜŞTERİ ADI SOYADI</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJE</th>
        """
        
        # Add day names
        for day_name in week_day_names:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center; background-color: #f39c12; color: white;">{day_name}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">GENEL TOPLAM</th></tr>'
        
        # Add date headers
        html += '<tr style="background-color: #fdf2e9;">'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th>'
        
        for date_str in week_dates:
            html += f'<th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">{date_str}</th>'
        
        html += '<th style="border: 1px solid #bdc3c7; padding: 8px;"></th></tr></thead><tbody>'
        
        # Empty row
        html += '<tr style="background-color: #fff8e1;">'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">-</td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px; text-align: center; color: #7f8c8d;">Bu hafta çek tahsilatı yok</td>'
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        for date_str in week_dates:
            html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td>'
        
        html += '<td style="border: 1px solid #bdc3c7; padding: 8px;"></td></tr>'
        html += '</tbody></table>'
        return html
    
    def _generate_simple_analysis_tables_html(self, payment_type_analysis: Dict, project_totals_analysis: Dict, 
                                            location_analysis: Dict, week_start: datetime, is_weekly_sheet: bool = True) -> str:
        """Generate simple HTML for analysis tables with side-by-side layout"""
        html = ""
        
        # Analysis Tables - Side by Side Layout using simple table structure
        html += """
            <div style="margin-top: 40px;">
                <table style="width: 100%; border-collapse: collapse;">
                    <tr>
                        <!-- Payment Type Analysis -->
                        <td style="width: 33.33%; vertical-align: top; padding: 10px;">
                            <div style="background-color: #f8f9fa; border-radius: 8px; padding: 15px;">
                                <h3 style="color: #8e44ad; border-bottom: 2px solid #8e44ad; padding-bottom: 5px; margin-top: 0; text-align: center;">
                                    ÖDEME TİPİ ANALİZİ
                                </h3>
                                
                                <table style="width: 100%; border-collapse: collapse;">
                                    <tr>
                                        <td style="width: 50%; vertical-align: top; padding-right: 5px;">
                                            <h4 style="color: #9b59b6; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Add weekly payment type analysis data
        weekly_analysis = payment_type_analysis.get('weekly', {})
        
        if is_weekly_sheet and week_start:
            # For weekly sheets, show data for the specific week
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_data = weekly_analysis.get(week_key, {
                'BANK_TRANSFER': {'tl_total': 0, 'usd_total': 0},
                'Nakit': {'tl_total': 0, 'usd_total': 0},
                'Çek': {'tl_total': 0, 'usd_total': 0},
                'Genel Toplam': {'tl_total': 0, 'usd_total': 0}
            })
        else:
            # For summary sheet, sum all weeks
            weekly_data = {
                'BANK_TRANSFER': {'tl_total': 0, 'usd_total': 0},
                'Nakit': {'tl_total': 0, 'usd_total': 0},
                'Çek': {'tl_total': 0, 'usd_total': 0},
                'Genel Toplam': {'tl_total': 0, 'usd_total': 0}
            }
            
            for week_key, week_data in weekly_analysis.items():
                for payment_type in ['BANK_TRANSFER', 'Nakit', 'Çek', 'Genel Toplam']:
                    if payment_type in week_data:
                        weekly_data[payment_type]['tl_total'] += week_data[payment_type].get('tl_total', 0)
                        weekly_data[payment_type]['usd_total'] += week_data[payment_type].get('usd_total', 0)
        
        html += """
                                            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                                                <thead>
                                                    <tr style="background-color: #f4f4f4;">
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Ödeme Nedeni</th>
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam TL</th>
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                                                    </tr>
                                                </thead>
                                                <tbody>
        """
        
        payment_types = ['BANK_TRANSFER', 'Nakit', 'Çek']
        payment_rows_rendered = 0
        payment_types_present = 0
        for payment_type in payment_types:
            data = weekly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
            if data['tl_total'] > 0 or data['usd_total'] > 0:
                html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{payment_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${data['usd_total']:,.2f}</td>
                    </tr>
                """
                payment_rows_rendered += 1
                payment_types_present += 1

        # Only show Ödenen Tutar row if more than one payment type is present in the week
        if is_weekly_sheet and payment_rows_rendered > 0 and payment_types_present > 1:
            paid_total_tl = sum(weekly_data[p]['tl_total'] for p in payment_types)
            paid_total_usd = sum(weekly_data[p]['usd_total'] for p in payment_types)
            html += f"""
                <tr style="background-color: #f9e79f; color: #2c3e50; font-weight: bold;">
                    <td style="border: 1px solid #bdc3c7; padding: 8px;">Ödenen Tutar</td>
                    <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{paid_total_tl:,.2f}</td>
                    <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${paid_total_usd:,.2f}</td>
                </tr>
            """
        
        # Total row
        total_data = weekly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
        html += f"""
                    <tr style="background-color: #8e44ad; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">Genel Toplam</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{total_data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_data['usd_total']:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                        <td style="width: 50%; vertical-align: top; padding-left: 5px;">
                                            <h4 style="color: #9b59b6; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Add monthly payment type analysis data
        monthly_data = payment_type_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Ödeme Nedeni</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam TL</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for payment_type in payment_types:
            data = monthly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{payment_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${data['usd_total']:,.2f}</td>
                    </tr>
            """
        
        # Monthly total row
        monthly_total = monthly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
        html += f"""
                    <tr style="background-color: #8e44ad; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">Genel Toplam</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{monthly_total['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${monthly_total['usd_total']:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                    </tr>
                                </table>
                            </div>
                        </td>
                        
                        <!-- Project Totals Analysis -->
                        <td style="width: 33.33%; vertical-align: top; padding: 10px;">
                            <div style="background-color: #f0f8ff; border-radius: 8px; padding: 15px;">
                                <h3 style="color: #2ecc71; border-bottom: 2px solid #2ecc71; padding-bottom: 5px; margin-top: 0; text-align: center;">
                                    PROJE TOPLAMLARI
                                </h3>
                                
                                <table style="width: 100%; border-collapse: collapse;">
                                    <tr>
                                        <td style="width: 50%; vertical-align: top; padding-right: 5px;">
                                            <h4 style="color: #2ecc71; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Add weekly project totals analysis data
        weekly_project_analysis = project_totals_analysis.get('weekly', {})
        
        if is_weekly_sheet and week_start:
            # For weekly sheets, show data for the specific week
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_project_data = weekly_project_analysis.get(week_key, {'PROJECT_A': 0, 'PROJECT_B': 0, 'TOPLAM': 0})
        else:
            # For summary sheet, sum all weeks
            weekly_project_data = {'PROJECT_A': 0, 'PROJECT_B': 0, 'TOPLAM': 0}
            for week_key, week_data in weekly_project_analysis.items():
                for project_type in ['PROJECT_A', 'PROJECT_B', 'TOPLAM']:
                    if project_type in week_data:
                        weekly_project_data[project_type] += week_data[project_type]
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Proje</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        project_types = ['PROJECT_A', 'PROJECT_B']
        for project_type in project_types:
            amount = weekly_project_data.get(project_type, 0)
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{project_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>
                    </tr>
            """
        
        # Total row
        total_amount = weekly_project_data.get('TOPLAM', 0)
        html += f"""
                    <tr style="background-color: #2ecc71; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_amount:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                        <td style="width: 50%; vertical-align: top; padding-left: 5px;">
                                            <h4 style="color: #2ecc71; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Add monthly project totals analysis data
        monthly_project_data = project_totals_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Proje</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for project_type in project_types:
            amount = monthly_project_data.get(project_type, 0)
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{project_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>
                    </tr>
            """
        
        # Monthly total row
        monthly_total_amount = monthly_project_data.get('TOPLAM', 0)
        html += f"""
                    <tr style="background-color: #2ecc71; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${monthly_total_amount:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                    </tr>
                                </table>
                            </div>
                        </td>
                        
                        <!-- Location Analysis -->
                        <td style="width: 33.33%; vertical-align: top; padding: 10px;">
                            <div style="background-color: #fff5f5; border-radius: 8px; padding: 15px;">
                                <h3 style="color: #e74c3c; border-bottom: 2px solid #e74c3c; padding-bottom: 5px; margin-top: 0; text-align: center;">
                                    LOKASYON ANALİZİ
                                </h3>
                                
                                <table style="width: 100%; border-collapse: collapse;">
                                    <tr>
                                        <td style="width: 50%; vertical-align: top; padding-right: 5px;">
                                            <h4 style="color: #c0392b; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Add weekly location analysis data
        weekly_location_analysis = location_analysis.get('weekly', {})
        
        if is_weekly_sheet and week_start:
            # For weekly sheets, show data for the specific week
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_location_data = weekly_location_analysis.get(week_key, {
                'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0}
            })
        else:
            # For summary sheet, sum all weeks
            weekly_location_data = {
                'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0}
            }
            
            for week_key, week_data in weekly_location_analysis.items():
                for location, project_data in week_data.items():
                    if location in weekly_location_data and isinstance(project_data, dict):
                        for project_type in ['PROJECT_A', 'PROJECT_B']:
                            if project_type in project_data:
                                weekly_location_data[location][project_type] += project_data[project_type]
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Lokasyon</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_A</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_B</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for location, project_data in weekly_location_data.items():
            PROJECT_A_amount = project_data.get('PROJECT_A', 0)
            PROJECT_B_amount = project_data.get('PROJECT_B', 0)
            total_amount = PROJECT_A_amount + PROJECT_B_amount
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{location}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_amount:,.2f}</td>
                    </tr>
            """
        
        # Calculate weekly totals
        PROJECT_A_total = sum(data.get('PROJECT_A', 0) for data in weekly_location_data.values() if isinstance(data, dict))
        PROJECT_B_total = sum(data.get('PROJECT_B', 0) for data in weekly_location_data.values() if isinstance(data, dict))
        general_total = PROJECT_A_total + PROJECT_B_total
        html += f"""
                    <tr style="background-color: #e74c3c; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${general_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                        <td style="width: 50%; vertical-align: top; padding-left: 5px;">
                                            <h4 style="color: #c0392b; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Add monthly location analysis data
        monthly_location_data = location_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Lokasyon</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_A</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_B</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for location, project_data in monthly_location_data.items():
            if isinstance(project_data, dict):
                PROJECT_A_amount = project_data.get('PROJECT_A', 0)
                PROJECT_B_amount = project_data.get('PROJECT_B', 0)
                total_amount = PROJECT_A_amount + PROJECT_B_amount
                html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{location}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_amount:,.2f}</td>
                    </tr>
            """
        
        # Calculate monthly totals
        PROJECT_A_total = sum(data.get('PROJECT_A', 0) for data in monthly_location_data.values() if isinstance(data, dict))
        PROJECT_B_total = sum(data.get('PROJECT_B', 0) for data in monthly_location_data.values() if isinstance(data, dict))
        general_total = PROJECT_A_total + PROJECT_B_total
        html += f"""
                    <tr style="background-color: #e74c3c; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${general_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                                        </td>
                                    </tr>
                                </table>
                            </div>
                        </td>
                    </tr>
                </table>
            </div>
        """
        
        return html
    
    def _generate_analysis_tables_html(self, payment_type_analysis: Dict, project_totals_analysis: Dict, 
                                     location_analysis: Dict, week_start: datetime) -> str:
        """Generate HTML for analysis tables with side-by-side layout"""
        html = ""
        
        # Analysis Tables - Side by Side Layout using table structure for better QTextEdit compatibility
        html += """
            <div style="margin-top: 40px;">
                <table style="width: 100%; border-collapse: collapse;">
                    <tr>
                        <!-- Payment Type Analysis -->
                        <td style="width: 33.33%; vertical-align: top; padding: 10px;">
                            <div style="background-color: #f8f9fa; border-radius: 8px; padding: 15px;">
                                <h3 style="color: #8e44ad; border-bottom: 2px solid #8e44ad; padding-bottom: 5px; margin-top: 0; text-align: center;">
                                    ÖDEME TİPİ ANALİZİ
                                </h3>
                                
                                <table style="width: 100%; border-collapse: collapse;">
                                    <tr>
                                        <td style="width: 50%; vertical-align: top; padding-right: 5px;">
                                            <h4 style="color: #9b59b6; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Weekly payment type analysis - sum all weeks in the analysis
        weekly_analysis = payment_type_analysis.get('weekly', {})
        weekly_data = {
            'BANK_TRANSFER': {'tl_total': 0, 'usd_total': 0},
            'Nakit': {'tl_total': 0, 'usd_total': 0},
            'Çek': {'tl_total': 0, 'usd_total': 0},
            'Genel Toplam': {'tl_total': 0, 'usd_total': 0}
        }
        
        # Sum all weeks
        for week_key, week_data in weekly_analysis.items():
            for payment_type in ['BANK_TRANSFER', 'Nakit', 'Çek', 'Genel Toplam']:
                if payment_type in week_data:
                    weekly_data[payment_type]['tl_total'] += week_data[payment_type].get('tl_total', 0)
                    weekly_data[payment_type]['usd_total'] += week_data[payment_type].get('usd_total', 0)
        
        html += """
                                            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                                                <thead>
                                                    <tr style="background-color: #f4f4f4;">
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Ödeme Nedeni</th>
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam TL</th>
                                                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                                                    </tr>
                                                </thead>
                                                <tbody>
        """
        
        payment_types = ['BANK_TRANSFER', 'Nakit', 'Çek']
        for payment_type in payment_types:
            data = weekly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{payment_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${data['usd_total']:,.2f}</td>
                    </tr>
            """
        
        # Total row
        total_data = weekly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
        html += f"""
                    <tr style="background-color: #8e44ad; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">Genel Toplam</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{total_data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_data['usd_total']:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                        </div>
                        
                        <div style="flex: 1;">
                            <h4 style="color: #9b59b6; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Monthly payment type analysis
        monthly_data = payment_type_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Ödeme Nedeni</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam TL</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for payment_type in payment_types:
            data = monthly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{payment_type}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{data['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${data['usd_total']:,.2f}</td>
                    </tr>
            """
        
        # Monthly total row
        monthly_total = monthly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
        html += f"""
                    <tr style="background-color: #8e44ad; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">Genel Toplam</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">₺{monthly_total['tl_total']:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${monthly_total['usd_total']:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                        </div>
                    </div>
                </div>
                
                <!-- Project Totals Analysis -->
                <div style="flex: 1; background-color: #f0f8ff; border-radius: 8px; padding: 15px; max-width: 48%;">
                    <h3 style="color: #2ecc71; border-bottom: 2px solid #2ecc71; padding-bottom: 5px; margin-top: 0; text-align: center;">
                        PROJE TOPLAMLARI
                    </h3>
                    
                    <div style="display: flex; gap: 10px; margin-bottom: 15px;">
                        <div style="flex: 1;">
                            <h4 style="color: #2ecc71; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Weekly project totals - sum all weeks in the analysis
        weekly_project_analysis = project_totals_analysis.get('weekly', {})
        weekly_project_data = {'PROJECT_A': 0, 'PROJECT_B': 0, 'TOPLAM': 0}
        
        # Sum all weeks
        for week_key, week_data in weekly_project_analysis.items():
            for project_type in ['PROJECT_A', 'PROJECT_B', 'TOPLAM']:
                if project_type in week_data:
                    weekly_project_data[project_type] += week_data[project_type]
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Proje</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        projects = ['PROJECT_A', 'PROJECT_B']
        for project in projects:
            amount = weekly_project_data.get(project, 0)
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{project}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>
                    </tr>
            """
        
        # Weekly total
        weekly_total = weekly_project_data.get('TOPLAM', 0)
        html += f"""
                    <tr style="background-color: #27ae60; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${weekly_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                        </div>
                        
                        <div style="flex: 1;">
                            <h4 style="color: #2ecc71; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Monthly project totals
        monthly_project_data = project_totals_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Proje</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for project in projects:
            amount = monthly_project_data.get(project, 0)
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{project}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${amount:,.2f}</td>
                    </tr>
            """
        
        # Monthly total
        monthly_total = monthly_project_data.get('TOPLAM', 0)
        html += f"""
                    <tr style="background-color: #27ae60; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${monthly_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                        </div>
                    </div>
                </div>
            </div>
            
            <!-- Location Analysis - Full Width -->
            <div style="background-color: #fff5f5; border-radius: 8px; padding: 20px; margin-top: 20px;">
                <h3 style="color: #e74c3c; border-bottom: 2px solid #e74c3c; padding-bottom: 5px; margin-top: 0; text-align: center;">
                    LOKASYON ANALİZİ
                </h3>
                
                <div style="display: flex; gap: 10px; margin-bottom: 15px;">
                    <div style="flex: 1;">
                        <h4 style="color: #c0392b; margin-bottom: 10px; font-size: 13px; text-align: center;">Haftalık</h4>
        """
        
        # Weekly location analysis - sum all weeks in the analysis
        weekly_location_analysis = location_analysis.get('weekly', {})
        weekly_location_data = {
            'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'TOPLAM': {'PROJECT_A': 0, 'PROJECT_B': 0}
        }
        
        # Sum all weeks
        for week_key, week_data in weekly_location_analysis.items():
            for location in weekly_location_data.keys():
                if location in week_data:
                    if isinstance(week_data[location], dict):
                        weekly_location_data[location]['PROJECT_A'] += week_data[location].get('PROJECT_A', 0)
                        weekly_location_data[location]['PROJECT_B'] += week_data[location].get('PROJECT_B', 0)
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Lokasyon</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_A USD</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_B USD</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        locations = ['ÇARŞI', 'LOCATION_B', 'OFİS', 'BANKA HAVALESİ', 'A KASA ÇEK', 'B KASA ÇEK']
        for location in locations:
            location_data = weekly_location_data.get(location, {'PROJECT_A': 0, 'PROJECT_B': 0})
            PROJECT_A_amount = location_data.get('PROJECT_A', 0) if isinstance(location_data, dict) else 0
            PROJECT_B_amount = location_data.get('PROJECT_B', 0) if isinstance(location_data, dict) else 0
            total_amount = PROJECT_A_amount + PROJECT_B_amount
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{location}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_amount:,.2f}</td>
                    </tr>
            """
        
        # Calculate weekly totals
        PROJECT_A_total = sum(data.get('PROJECT_A', 0) for data in weekly_location_data.values() if isinstance(data, dict))
        PROJECT_B_total = sum(data.get('PROJECT_B', 0) for data in weekly_location_data.values() if isinstance(data, dict))
        general_total = PROJECT_A_total + PROJECT_B_total
        html += f"""
                    <tr style="background-color: #e74c3c; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${general_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                    </div>
                    
                    <div style="flex: 1;">
                        <h4 style="color: #c0392b; margin-bottom: 10px; font-size: 13px; text-align: center;">Aylık</h4>
        """
        
        # Monthly location analysis
        monthly_location_data = location_analysis.get('monthly', {})
        
        html += """
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <thead>
                    <tr style="background-color: #f4f4f4;">
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Lokasyon</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_A USD</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">PROJECT_B USD</th>
                        <th style="border: 1px solid #bdc3c7; padding: 8px; text-align: center;">Toplam USD</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for location in locations:
            location_data = monthly_location_data.get(location, {'PROJECT_A': 0, 'PROJECT_B': 0})
            PROJECT_A_amount = location_data.get('PROJECT_A', 0) if isinstance(location_data, dict) else 0
            PROJECT_B_amount = location_data.get('PROJECT_B', 0) if isinstance(location_data, dict) else 0
            total_amount = PROJECT_A_amount + PROJECT_B_amount
            html += f"""
                    <tr style="background-color: white;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">{location}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_amount:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${total_amount:,.2f}</td>
                    </tr>
            """
        
        # Calculate monthly totals
        PROJECT_A_total = sum(data.get('PROJECT_A', 0) for data in monthly_location_data.values() if isinstance(data, dict))
        PROJECT_B_total = sum(data.get('PROJECT_B', 0) for data in monthly_location_data.values() if isinstance(data, dict))
        general_total = PROJECT_A_total + PROJECT_B_total
        html += f"""
                    <tr style="background-color: #e74c3c; color: white; font-weight: bold;">
                        <td style="border: 1px solid #bdc3c7; padding: 8px;">TOPLAM</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_A_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${PROJECT_B_total:,.2f}</td>
                        <td style="border: 1px solid #bdc3c7; padding: 8px; text-align: right;">${general_total:,.2f}</td>
                    </tr>
        """
        
        html += """
                </tbody>
            </table>
                    </div>
                </div>
            </div>
        """
        
        return html
    
    def _generate_summary_html(self, payments: List[PaymentData], start_date: datetime, end_date: datetime) -> str:
        """Generate HTML summary with analysis tables"""
        total_payments = len(payments)
        total_amount = sum(p.amount for p in payments)
        check_payments = [p for p in payments if hasattr(p, 'is_check_payment') and p.is_check_payment]
        total_check_amount = sum(p.cek_tutari if hasattr(p, 'cek_tutari') else p.amount for p in check_payments)
        
        # Generate analysis data
        payment_type_analysis = self.generate_payment_type_analysis(payments, start_date, end_date)
        project_totals_analysis = self.generate_project_totals_analysis(payments, start_date, end_date)
        location_analysis = self.generate_location_analysis(payments, start_date, end_date)
        
        html = f"""
        <div style="font-family: Arial, sans-serif; padding: 20px;">
            <h1 style="text-align: center; color: #2c3e50;">Rapor Özeti</h1>
            <div style="background-color: #ecf0f1; padding: 20px; border-radius: 8px; margin: 20px 0;">
                <h3 style="color: #2980b9;">Genel İstatistikler</h3>
                <p><strong>Tarih Aralığı:</strong> {start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}</p>
                <p><strong>Toplam Ödeme Sayısı:</strong> {total_payments}</p>
                <p><strong>Toplam Tutar:</strong> ${total_amount:,.2f}</p>
                <p><strong>Çek Ödemesi Sayısı:</strong> {len(check_payments)}</p>
                <p><strong>Toplam Çek Tutarı:</strong> ₺{total_check_amount:,.2f}</p>
            </div>
        """
        
        # Add analysis tables
        html += self._generate_simple_analysis_tables_html(payment_type_analysis, project_totals_analysis, location_analysis, start_date, is_weekly_sheet=False)
        
        html += """
        </div>
        """
        
        return html
    
    def export_to_excel(self, payments: List[PaymentData], 
                       start_date: datetime, end_date: datetime, 
                       output_path: str) -> None:
        """Export all reports to a single Excel file with Turkish character support"""
        try:
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                workbook = writer.book
                
                # Define formats with Turkish character support
                header_format = workbook.add_format({
                    'bold': True,
                    'text_wrap': True,
                    'valign': 'top',
                    'fg_color': '#D7E4BC',
                    'border': 1,
                    'font_name': 'Arial'
                })
                
                title_format = workbook.add_format({
                    'bold': True,
                    'font_size': 14,
                    'align': 'center',
                    'valign': 'vcenter',
                    'fg_color': '#4F81BD',
                    'font_color': 'white',
                    'border': 1,
                    'font_name': 'Arial'
                })
                
                subtitle_format = workbook.add_format({
                    'bold': True,
                    'font_size': 12,
                    'align': 'center',
                    'valign': 'vcenter',
                    'fg_color': '#B1C5E7',
                    'border': 1,
                    'font_name': 'Arial'
                })
                
                number_format = workbook.add_format({
                    'num_format': '#,##0.00',
                    'font_name': 'Arial'
                })
                currency_format = workbook.add_format({
                    'num_format': '$#,##0.00',
                    'font_name': 'Arial'
                })
                
                # Default cell format for Turkish characters
                default_format = workbook.add_format({
                    'font_name': 'Arial'
                })
                
                # Generate reports
                daily_breakdown = self.generate_daily_usd_breakdown(payments, start_date, end_date)
                customer_date_table = self.generate_customer_date_table(payments, start_date, end_date)
                customer_check_table = self.generate_customer_check_table(payments, start_date, end_date)
                weekly_summary = self.generate_weekly_summary(payments)
                monthly_summary = self.generate_monthly_summary(payments)
                daily_timeline = self.generate_daily_timeline(payments, start_date, end_date)
                payment_type_summary = self.generate_payment_type_summary(payments)
                
                # Create main customer-date table sheet first
                if customer_date_table:
                    # Create main sheet with complete customer-date table
                    main_worksheet = workbook.add_worksheet('Müşteri Tarih Tablosu')
                    
                    # Get all unique customers and projects across all weeks
                    all_customers = set()
                    all_dates = set()
                    for week_data in customer_date_table.values():
                        pivot = week_data['pivot']
                        if not pivot.empty:
                            all_customers.update(pivot.index.get_level_values(0))
                            all_dates.update([col for col in pivot.columns if col != 'Genel Toplam'])
                    
                    if all_customers and all_dates:
                        # Sort customers alphabetically
                        sorted_customers = sorted(all_customers)
                        sorted_dates = sorted(all_dates)
                        
                        # Calculate total columns needed
                        total_cols = 3 + len(sorted_dates) + 1  # SIRA NO + Müşteri + Proje + dates + Total
                        last_col_letter = chr(ord('A') + total_cols - 1)
                        
                        # Write main title
                        title_text = "MODEL KUYUM MERKEZİ - MODEL SANAYİ MERKEZİ TARİHLER TABLOSU"
                        main_worksheet.merge_range(f'A1:{last_col_letter}1', title_text, title_format)
                        
                        # Write subtitle
                        date_range_str = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
                        subtitle_text = f"{date_range_str} | BANK_TRANSFER - Nakit"
                        main_worksheet.merge_range(f'A2:{last_col_letter}2', subtitle_text, subtitle_format)
                        
                        # Write headers
                        current_row = 3
                        main_worksheet.write(current_row, 0, 'SIRA NO', header_format)
                        main_worksheet.write(current_row, 1, 'MÜŞTERİ ADI SOYADI', header_format)
                        main_worksheet.write(current_row, 2, 'PROJE', header_format)
                        
                        # Write date headers
                        col_idx = 3
                        for date_str in sorted_dates:
                            main_worksheet.write(current_row, col_idx, date_str, header_format)
                            col_idx += 1
                        
                        main_worksheet.write(current_row, col_idx, 'GENEL TOPLAM', header_format)
                        current_row += 1
                        
                        # Write data rows
                        sira_no = 1
                        for customer in sorted_customers:
                            # Get all projects for this customer
                            customer_projects = set()
                            for week_data in customer_date_table.values():
                                pivot = week_data['pivot']
                                if not pivot.empty:
                                    customer_projects.update(pivot.loc[customer].index if customer in pivot.index else [])
                            
                            for project in sorted(customer_projects):
                                main_worksheet.write(current_row, 0, sira_no)
                                main_worksheet.write(current_row, 1, customer)
                                main_worksheet.write(current_row, 2, project)
                                
                                # Calculate totals for this customer-project combination
                                total_amount = 0
                                col_idx = 3
                                
                                for date_str in sorted_dates:
                                    amount = 0
                                    # Sum amounts across all weeks for this customer-project-date
                                    for week_data in customer_date_table.values():
                                        pivot = week_data['pivot']
                                        if not pivot.empty and (customer, project) in pivot.index and date_str in pivot.columns:
                                            amount += pivot.loc[(customer, project), date_str]
                                    
                                    if amount > 0:
                                        main_worksheet.write(current_row, col_idx, amount, currency_format)
                                        total_amount += amount
                                    else:
                                        main_worksheet.write(current_row, col_idx, '')
                                    
                                    col_idx += 1
                                
                                # Write total
                                if total_amount > 0:
                                    main_worksheet.write(current_row, col_idx, total_amount, currency_format)
                                else:
                                    main_worksheet.write(current_row, col_idx, '')
                                
                                current_row += 1
                                sira_no += 1
                        
                        # Set column widths
                        main_worksheet.set_column(0, 0, 8)   # SIRA NO
                        main_worksheet.set_column(1, 1, 25)  # Müşteri
                        main_worksheet.set_column(2, 2, 20)  # Proje
                        for i in range(3, 3 + len(sorted_dates)):
                            main_worksheet.set_column(i, i, 12)  # Date columns
                        main_worksheet.set_column(3 + len(sorted_dates), 3 + len(sorted_dates), 15)  # Total
                
                # Create separate sheets for each week
                if customer_date_table:
                    sorted_weeks = sorted(customer_date_table.keys())
                    
                    # Define TL conversion highlight format
                    tl_highlight_format = workbook.add_format({
                        'num_format': '$#,##0.00',
                        'bg_color': '#E6F3FF',  # Light blue background
                        'border': 1
                    })
                    
                    for week_idx, week_start in enumerate(sorted_weeks):
                        week_data = customer_date_table[week_start]
                        pivot = week_data['pivot']
                        tl_converted = week_data['tl_converted']
                        conversion_rates = week_data['conversion_rates']
                        day_names = week_data['day_names']
                        week_range = week_data['week_range']
                        week_dates = week_data['week_dates']
                        week_day_names = week_data['week_day_names']
                        
                        if pivot.empty:
                            continue
                        
                        # Create sheet for this week
                        sheet_name = f'Hafta {week_idx + 1}'
                        worksheet = workbook.add_worksheet(sheet_name)
                        
                        # Calculate total columns needed: SIRA NO + Müşteri + Proje + 7 days + Total = 11
                        total_cols = 11
                        last_col_letter = chr(ord('A') + total_cols - 1)  # K
                        
                        # Write main title (properly merged across actual columns)
                        title_text = "MODEL KUYUM MERKEZİ - MODEL SANAYİ MERKEZİ TARİHLER TABLOSU"
                        worksheet.merge_range(f'A1:{last_col_letter}1', title_text, title_format)
                        
                        # Write subtitle with date range and payment type
                        date_range_str = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
                        subtitle_text = f"{date_range_str} | BANK_TRANSFER - Nakit"
                        worksheet.merge_range(f'A2:{last_col_letter}2', subtitle_text, subtitle_format)
                        
                        # Write week header
                        week_header_text = f"HAFTA {week_idx + 1}: {week_range}"
                        worksheet.merge_range(f'A3:{last_col_letter}3', week_header_text, subtitle_format)
                        
                        # Write Turkish day names row
                        current_row = 4
                        worksheet.write(current_row, 0, '', header_format)  # SIRA NO column
                        worksheet.write(current_row, 1, '', header_format)  # Müşteri column
                        worksheet.write(current_row, 2, '', header_format)  # Proje column
                        
                        # Write Turkish day names for all 7 days (remove HTML tags for Excel)
                        col_idx = 3
                        for day_name in week_day_names:
                            # Remove HTML tags for Excel export
                            clean_day_name = day_name.replace('<br>', '\n')
                            worksheet.write(current_row, col_idx, clean_day_name, header_format)
                            col_idx += 1
                        
                        # Fill remaining days if less than 7
                        while col_idx < 10:  # Up to column J (7 days)
                            worksheet.write(current_row, col_idx, '', header_format)
                            col_idx += 1
                        
                        worksheet.write(current_row, 10, '', header_format)  # Genel Toplam column
                        current_row += 1
                        
                        # Write date headers
                        worksheet.write(current_row, 0, 'SIRA NO', header_format)
                        worksheet.write(current_row, 1, 'MÜŞTERİ ADI SOYADI', header_format)
                        worksheet.write(current_row, 2, 'PROJE', header_format)
                        
                        # Write dates for the week (all 7 days)
                        col_idx = 3
                        for date_str in week_dates:
                            worksheet.write(current_row, col_idx, date_str, header_format)
                            col_idx += 1
                        
                        # Fill remaining date columns if less than 7 days
                        while col_idx < 10:
                            worksheet.write(current_row, col_idx, '', header_format)
                            col_idx += 1
                        
                        worksheet.write(current_row, 10, 'GENEL TOPLAM', header_format)
                        current_row += 1
                        
                        # Write data rows with SIRA NO
                        sira_no = 1
                        for (customer, project), row_data in pivot.iterrows():
                            worksheet.write(current_row, 0, sira_no)  # SIRA NO
                            worksheet.write(current_row, 1, customer)  # Customer
                            worksheet.write(current_row, 2, project)   # Project
                            
                            # Write amounts for each date (7 days)
                            col_idx = 3
                            for date_str in week_dates:
                                amount = row_data[date_str] if date_str in row_data.index else 0
                                is_tl = tl_converted.loc[(customer, project), date_str] if date_str in tl_converted.columns else False
                                
                                if amount > 0:
                                    # Use highlight format for TL converted amounts
                                    format_to_use = tl_highlight_format if is_tl else currency_format
                                    worksheet.write(current_row, col_idx, amount, format_to_use)
                                else:
                                    worksheet.write(current_row, col_idx, '')
                                col_idx += 1
                            
                            # Fill remaining columns if less than 7 days
                            while col_idx < 10:
                                worksheet.write(current_row, col_idx, '')
                                col_idx += 1
                            
                            # Write total
                            total_amount = row_data['Genel Toplam']
                            if total_amount > 0:
                                worksheet.write(current_row, 10, total_amount, currency_format)
                            else:
                                worksheet.write(current_row, 10, '')
                            
                            current_row += 1
                            sira_no += 1
                        
                        # Add summary row for this week
                        worksheet.write(current_row, 0, '', header_format)
                        worksheet.write(current_row, 1, 'HAFTA TOPLAMI', header_format)
                        worksheet.write(current_row, 2, '', header_format)
                        
                        # Calculate column totals for 7 days
                        col_idx = 3
                        for date_str in week_dates:
                            total = pivot[date_str].sum() if date_str in pivot.columns else 0
                            if total > 0:
                                worksheet.write(current_row, col_idx, total, currency_format)
                            else:
                                worksheet.write(current_row, col_idx, '', header_format)
                            col_idx += 1
                        
                        # Fill remaining total columns
                        while col_idx < 10:
                            worksheet.write(current_row, col_idx, '', header_format)
                            col_idx += 1
                        
                        # Week total
                        week_total = pivot['Genel Toplam'].sum()
                        if week_total > 0:
                            worksheet.write(current_row, 10, week_total, currency_format)
                        else:
                            worksheet.write(current_row, 10, '', header_format)
                        
                        # Add check table if there are check payments for this week
                        if customer_check_table and week_start in customer_check_table:
                            check_data = customer_check_table[week_start]
                            check_pivot_tl = check_data['pivot_tl']
                            check_pivot_usd = check_data['pivot_usd']
                            
                            if not check_pivot_tl.empty:
                                current_row += 3  # Add some spacing
                                
                                # Check table title
                                check_title = f"ÇEK TAHSİLATLARI - {week_range}"
                                worksheet.merge_range(f'A{current_row}:{last_col_letter}{current_row}', check_title, subtitle_format)
                                current_row += 1
                                
                                # Turkish day names row for checks
                                worksheet.write(current_row, 0, '', header_format)  # SIRA NO column
                                worksheet.write(current_row, 1, '', header_format)  # Müşteri column
                                worksheet.write(current_row, 2, '', header_format)  # Proje column
                                
                                col_idx = 3
                                for day_name in week_day_names:
                                    worksheet.write(current_row, col_idx, day_name, header_format)
                                    col_idx += 1
                                
                                while col_idx < 10:
                                    worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                worksheet.write(current_row, 10, '', header_format)
                                current_row += 1
                                
                                # Date headers for checks
                                worksheet.write(current_row, 0, 'SIRA NO', header_format)
                                worksheet.write(current_row, 1, 'MÜŞTERİ ADI SOYADI', header_format)
                                worksheet.write(current_row, 2, 'PROJE', header_format)
                                
                                col_idx = 3
                                for date_str in week_dates:
                                    worksheet.write(current_row, col_idx, date_str, header_format)
                                    col_idx += 1
                                
                                while col_idx < 10:
                                    worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                worksheet.write(current_row, 10, 'GENEL TOPLAM', header_format)
                                current_row += 1
                                
                                # TL Check amounts - same structure as regular payments
                                sira_no = 1
                                for (customer, project), row_data in check_pivot_tl.iterrows():
                                    worksheet.write(current_row, 0, sira_no)
                                    worksheet.write(current_row, 1, customer)
                                    worksheet.write(current_row, 2, project)
                                    
                                    col_idx = 3
                                    for date_str in week_dates:
                                        amount = row_data.get(date_str, 0) if hasattr(row_data, 'get') else 0
                                        if amount > 0:
                                            worksheet.write(current_row, col_idx, amount, number_format)
                                        else:
                                            worksheet.write(current_row, col_idx, '')
                                        col_idx += 1
                                    
                                    while col_idx < 10:
                                        worksheet.write(current_row, col_idx, '')
                                        col_idx += 1
                                    
                                    total_amount = row_data.get('Genel Toplam', 0) if hasattr(row_data, 'get') else 0
                                    if total_amount > 0:
                                        worksheet.write(current_row, 10, total_amount, number_format)
                                    else:
                                        worksheet.write(current_row, 10, '')
                                    
                                    current_row += 1
                                    sira_no += 1
                                
                                # TL Total row
                                worksheet.write(current_row, 0, '', header_format)
                                worksheet.write(current_row, 1, 'TOPLAM TL', header_format)
                                worksheet.write(current_row, 2, '', header_format)
                                
                                col_idx = 3
                                for date_str in week_dates:
                                    total = check_pivot_tl[date_str].sum() if date_str in check_pivot_tl.columns else 0
                                    if total > 0:
                                        worksheet.write(current_row, col_idx, total, number_format)
                                    else:
                                        worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                while col_idx < 10:
                                    worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                tl_total = check_pivot_tl['Genel Toplam'].sum()
                                if tl_total > 0:
                                    worksheet.write(current_row, 10, tl_total, number_format)
                                else:
                                    worksheet.write(current_row, 10, '', header_format)
                                
                                current_row += 1
                                
                                # USD Check amounts (using maturity date rates)
                                sira_no = 1
                                for (customer, project), row_data in check_pivot_usd.iterrows():
                                    worksheet.write(current_row, 0, sira_no)
                                    worksheet.write(current_row, 1, customer)
                                    worksheet.write(current_row, 2, project)
                                    
                                    col_idx = 3
                                    for date_str in week_dates:
                                        amount = row_data.get(date_str, 0) if hasattr(row_data, 'get') else 0
                                        if amount > 0:
                                            worksheet.write(current_row, col_idx, amount, currency_format)
                                        else:
                                            worksheet.write(current_row, col_idx, '')
                                        col_idx += 1
                                    
                                    while col_idx < 10:
                                        worksheet.write(current_row, col_idx, '')
                                        col_idx += 1
                                    
                                    total_amount = row_data.get('Genel Toplam', 0) if hasattr(row_data, 'get') else 0
                                    if total_amount > 0:
                                        worksheet.write(current_row, 10, total_amount, currency_format)
                                    else:
                                        worksheet.write(current_row, 10, '')
                                    
                                    current_row += 1
                                    sira_no += 1
                                
                                # USD Total row
                                worksheet.write(current_row, 0, '', header_format)
                                worksheet.write(current_row, 1, 'TOPLAM USD (Vade Tarihi Kuru)', header_format)
                                worksheet.write(current_row, 2, '', header_format)
                                
                                col_idx = 3
                                for date_str in week_dates:
                                    total = check_pivot_usd[date_str].sum() if date_str in check_pivot_usd.columns else 0
                                    if total > 0:
                                        worksheet.write(current_row, col_idx, total, currency_format)
                                    else:
                                        worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                while col_idx < 10:
                                    worksheet.write(current_row, col_idx, '', header_format)
                                    col_idx += 1
                                
                                usd_total = check_pivot_usd['Genel Toplam'].sum()
                                if usd_total > 0:
                                    worksheet.write(current_row, 10, usd_total, currency_format)
                                else:
                                    worksheet.write(current_row, 10, '', header_format)
                        else:
                            # Add empty check table if no checks for this week
                            current_row += 3
                            check_title = f"ÇEK TAHSİLATLARI - {week_range}"
                            worksheet.merge_range(f'A{current_row}:{last_col_letter}{current_row}', check_title, subtitle_format)
                            current_row += 1
                            
                            # Empty check table structure
                            worksheet.write(current_row, 0, '', header_format)
                            worksheet.write(current_row, 1, 'Bu hafta çek tahsilatı yok', header_format)
                            worksheet.write(current_row, 2, '', header_format)
                            for col_idx in range(3, 11):
                                worksheet.write(current_row, col_idx, '', header_format)
                            current_row += 1
                        
                        # Add analysis tables for this week
                        current_row += 3
                        
                        # Generate analysis data for this week
                        week_end = week_start + timedelta(days=6)
                        if week_end > end_date:
                            week_end = end_date
                        
                        payment_type_analysis = self.generate_payment_type_analysis(payments, week_start, week_end)
                        project_totals_analysis = self.generate_project_totals_analysis(payments, week_start, week_end)
                        location_analysis = self.generate_location_analysis(payments, week_start, week_end)
                        
                        # Payment Type Analysis Table
                        current_row += 1
                        payment_type_title = "ÖDEME TİPİ ANALİZİ"
                        worksheet.merge_range(f'A{current_row}:{last_col_letter}{current_row}', payment_type_title, subtitle_format)
                        current_row += 1
                        
                        # Weekly payment type analysis
                        week_key = week_start.strftime('%Y-%m-%d')
                        weekly_data = payment_type_analysis.get('weekly', {}).get(week_key, {})
                        
                        # Headers
                        worksheet.write(current_row, 0, 'Ödeme Nedeni', header_format)
                        worksheet.write(current_row, 1, 'Toplam TL', header_format)
                        worksheet.write(current_row, 2, 'Toplam USD', header_format)
                        worksheet.write(current_row, 3, 'Ödeme Nedeni', header_format)
                        worksheet.write(current_row, 4, 'Toplam TL', header_format)
                        worksheet.write(current_row, 5, 'Toplam USD', header_format)
                        current_row += 1
                        
                        # Weekly data
                        payment_types = ['BANK_TRANSFER', 'Nakit', 'Çek']
                        for i, payment_type in enumerate(payment_types):
                            data = weekly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
                            worksheet.write(current_row, 0, payment_type)
                            worksheet.write(current_row, 1, data['tl_total'], number_format)
                            worksheet.write(current_row, 2, data['usd_total'], currency_format)
                            
                            # Monthly data
                            monthly_data = payment_type_analysis.get('monthly', {})
                            monthly_payment_data = monthly_data.get(payment_type, {'tl_total': 0, 'usd_total': 0})
                            worksheet.write(current_row, 3, payment_type)
                            worksheet.write(current_row, 4, monthly_payment_data['tl_total'], number_format)
                            worksheet.write(current_row, 5, monthly_payment_data['usd_total'], currency_format)
                            current_row += 1
                        
                        # Totals
                        total_data = weekly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
                        monthly_total = monthly_data.get('Genel Toplam', {'tl_total': 0, 'usd_total': 0})
                        worksheet.write(current_row, 0, 'Genel Toplam', header_format)
                        worksheet.write(current_row, 1, total_data['tl_total'], number_format)
                        worksheet.write(current_row, 2, total_data['usd_total'], currency_format)
                        worksheet.write(current_row, 3, 'Genel Toplam', header_format)
                        worksheet.write(current_row, 4, monthly_total['tl_total'], number_format)
                        worksheet.write(current_row, 5, monthly_total['usd_total'], currency_format)
                        current_row += 2
                        
                        # Project Totals Analysis Table
                        project_title = "PROJE TOPLAMLARI"
                        worksheet.merge_range(f'A{current_row}:{last_col_letter}{current_row}', project_title, subtitle_format)
                        current_row += 1
                        
                        # Headers
                        worksheet.write(current_row, 0, 'Proje', header_format)
                        worksheet.write(current_row, 1, 'Toplam USD', header_format)
                        worksheet.write(current_row, 2, 'Proje', header_format)
                        worksheet.write(current_row, 3, 'Toplam USD', header_format)
                        current_row += 1
                        
                        # Weekly project data - sum all weeks
                        weekly_project_analysis = project_totals_analysis.get('weekly', {})
                        weekly_project_data = {'PROJECT_A': 0, 'PROJECT_B': 0, 'TOPLAM': 0}
                        for wk_key, wk_data in weekly_project_analysis.items():
                            for project_type in ['PROJECT_A', 'PROJECT_B', 'TOPLAM']:
                                if project_type in wk_data:
                                    weekly_project_data[project_type] += wk_data[project_type]
                        monthly_project_data = project_totals_analysis.get('monthly', {})
                        
                        projects = ['PROJECT_A', 'PROJECT_B']
                        for i, project in enumerate(projects):
                            weekly_amount = weekly_project_data.get(project, 0)
                            monthly_amount = monthly_project_data.get(project, 0)
                            worksheet.write(current_row, 0, project)
                            worksheet.write(current_row, 1, weekly_amount, currency_format)
                            worksheet.write(current_row, 2, project)
                            worksheet.write(current_row, 3, monthly_amount, currency_format)
                            current_row += 1
                        
                        # Project totals
                        weekly_total = weekly_project_data.get('TOPLAM', 0)
                        monthly_total = monthly_project_data.get('TOPLAM', 0)
                        worksheet.write(current_row, 0, 'TOPLAM', header_format)
                        worksheet.write(current_row, 1, weekly_total, currency_format)
                        worksheet.write(current_row, 2, 'TOPLAM', header_format)
                        worksheet.write(current_row, 3, monthly_total, currency_format)
                        current_row += 2
                        
                        # Location Analysis Table
                        location_title = "LOKASYON ANALİZİ"
                        worksheet.merge_range(f'A{current_row}:{last_col_letter}{current_row}', location_title, subtitle_format)
                        current_row += 1
                        
                        # Headers
                        worksheet.write(current_row, 0, 'Lokasyon', header_format)
                        worksheet.write(current_row, 1, 'Toplam USD', header_format)
                        worksheet.write(current_row, 2, 'Lokasyon', header_format)
                        worksheet.write(current_row, 3, 'Toplam USD', header_format)
                        current_row += 1
                        
                        # Weekly location data - sum all weeks
                        weekly_location_analysis = location_analysis.get('weekly', {})
                        weekly_location_data = {
                            'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
                            'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
                            'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
                            'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
                            'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
                            'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0}
                        }
                        for wk_key, wk_data in weekly_location_analysis.items():
                            for location in weekly_location_data.keys():
                                if location in wk_data and isinstance(wk_data[location], dict):
                                    weekly_location_data[location]['PROJECT_A'] += wk_data[location].get('PROJECT_A', 0)
                                    weekly_location_data[location]['PROJECT_B'] += wk_data[location].get('PROJECT_B', 0)
                        monthly_location_data = location_analysis.get('monthly', {})
                        
                        locations = ['ÇARŞI', 'LOCATION_B', 'OFİS', 'BANKA HAVALESİ', 'A KASA ÇEK', 'B KASA ÇEK']
                        for i, location in enumerate(locations):
                            # Handle dictionary structure from location analysis
                            weekly_data = weekly_location_data.get(location, {})
                            monthly_data = monthly_location_data.get(location, {})
                            
                            if isinstance(weekly_data, dict):
                                weekly_amount = weekly_data.get('PROJECT_A', 0) + weekly_data.get('PROJECT_B', 0)
                            else:
                                weekly_amount = weekly_data
                                
                            if isinstance(monthly_data, dict):
                                monthly_amount = monthly_data.get('PROJECT_A', 0) + monthly_data.get('PROJECT_B', 0)
                            else:
                                monthly_amount = monthly_data
                            
                            worksheet.write(current_row, 0, location)
                            worksheet.write(current_row, 1, weekly_amount, currency_format)
                            worksheet.write(current_row, 2, location)
                            worksheet.write(current_row, 3, monthly_amount, currency_format)
                            current_row += 1
                        
                        # Location totals
                        weekly_total_data = weekly_location_data.get('TOPLAM', {})
                        monthly_total_data = monthly_location_data.get('TOPLAM', {})
                        
                        if isinstance(weekly_total_data, dict):
                            weekly_location_total = weekly_total_data.get('GENEL', 0)
                        else:
                            weekly_location_total = weekly_total_data
                            
                        if isinstance(monthly_total_data, dict):
                            monthly_location_total = monthly_total_data.get('GENEL', 0)
                        else:
                            monthly_location_total = monthly_total_data
                        worksheet.write(current_row, 0, 'TOPLAM', header_format)
                        worksheet.write(current_row, 1, weekly_location_total, currency_format)
                        worksheet.write(current_row, 2, 'TOPLAM', header_format)
                        worksheet.write(current_row, 3, monthly_location_total, currency_format)
                        
                        # Set column widths
                        worksheet.set_column('A:A', 8)   # SIRA NO
                        worksheet.set_column('B:B', 25)  # Customer name
                        worksheet.set_column('C:C', 15)  # Project
                        worksheet.set_column('D:J', 12)  # Date columns (7 days)
                        worksheet.set_column('K:K', 15)  # Total column
                
                # Write other reports to separate sheets
                if not daily_breakdown.empty:
                    daily_breakdown.to_excel(writer, sheet_name='Günlük USD Dağılımı')
                    worksheet = writer.sheets['Günlük USD Dağılımı']
                    worksheet.set_column('A:A', 20)  # Customer column
                    worksheet.set_column('B:B', 25)  # Project column
                    for col in range(2, len(daily_breakdown.columns) + 2):
                        worksheet.set_column(col, col, 15, currency_format)
                
                if not weekly_summary.empty:
                    weekly_summary.to_excel(writer, sheet_name='Haftalık Özet')
                    worksheet = writer.sheets['Haftalık Özet']
                    worksheet.set_column('A:A', 25)  # Project column
                    for col in range(1, len(weekly_summary.columns) + 1):
                        worksheet.set_column(col, col, 15, currency_format)
                
                if not monthly_summary.empty:
                    monthly_summary.to_excel(writer, sheet_name='Aylık Kanal Dağılımı')
                    worksheet = writer.sheets['Aylık Kanal Dağılımı']
                    worksheet.set_column('A:A', 20)  # Channel column
                    for col in range(1, len(monthly_summary.columns) + 1):
                        worksheet.set_column(col, col, 15, currency_format)
                
                if not daily_timeline.empty:
                    daily_timeline.to_excel(writer, sheet_name='Günlük Zaman Çizelgesi')
                    worksheet = writer.sheets['Günlük Zaman Çizelgesi']
                    worksheet.set_column('A:A', 12)  # Date column
                    for col in range(1, len(daily_timeline.columns) + 1):
                        worksheet.set_column(col, col, 15, currency_format)
                
                if not payment_type_summary.empty:
                    payment_type_summary.to_excel(writer, sheet_name='Ödeme Türü Özeti')
                    worksheet = writer.sheets['Ödeme Türü Özeti']
                    worksheet.set_column('A:A', 20)  # Channel column
                    worksheet.set_column('B:B', 15, number_format)  # TL column
                    worksheet.set_column('C:C', 15, currency_format)  # USD column
            
            logger.info(f"Excel report exported to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export Excel report: {e}")
            raise
    
    def _export_single_week_to_excel(self, workbook, sheet_name, week_data, check_data, payments, start_date, end_date):
        """Export a single week's data to Excel with professional formatting matching the UI"""
        try:
            worksheet = workbook.add_worksheet(sheet_name)
            
            # Define professional styles matching UI appearance
            # Main title format (matches UI's large, bold title)
            main_title_format = workbook.add_format({
                'bold': True,
                'font_size': 18,
                'font_color': '#2C3E50',
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#F8F9FA',
                'border': 1,
                'text_wrap': True
            })
            
            # Week title format (matches UI's week title styling)
            week_title_format = workbook.add_format({
                'bold': True,
                'font_size': 16,
                'font_color': '#2C3E50',
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#F8F9FA',
                'border': 1,
                'text_wrap': True
            })
            
            # Section header format (matches UI's table headers)
            section_header_format = workbook.add_format({
                'bold': True,
                'font_size': 14,
                'font_color': '#495057',
                'bg_color': '#F8F9FA',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'text_wrap': True
            })
            
            # Table header format (matches UI's QHeaderView styling)
            table_header_format = workbook.add_format({
                'bold': True,
                'font_weight': 600,
                'font_color': '#495057',
                'bg_color': '#F8F9FA',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'text_wrap': True
            })
            
            # Data cell format (matches UI's table cells)
            data_cell_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'font_color': '#212529'
            })
            
            # Customer name format (left-aligned for better readability)
            customer_cell_format = workbook.add_format({
                'border': 1,
                'align': 'left',
                'valign': 'vcenter',
                'font_color': '#212529'
            })
            
            # Currency format with Turkish Lira symbol
            currency_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'num_format': '₺#,##0.00',
                'font_color': '#212529'
            })
            
            # Total row format (highlighted)
            total_row_format = workbook.add_format({
                'bold': True,
                'bg_color': '#E9ECEF',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'font_color': '#212529'
            })
            
            current_row = 0
            
            # Main title with professional styling
            worksheet.merge_range(current_row, 0, current_row, 10, 
                                "MODEL KUYUM MERKEZİ – MODEL SANAYİ MERKEZİ TARİHLER TABLOSU", 
                                main_title_format)
            current_row += 1
            
            # Week range with improved styling
            week_range = week_data.get('week_range', '')
            worksheet.merge_range(current_row, 0, current_row, 10, 
                                f"{week_range}", week_title_format)
            current_row += 2
            
            # Normal Payments Section with better visual hierarchy
            worksheet.merge_range(current_row, 0, current_row, 10, 
                                "HAFTANIN TÜM ÖDEMELERİ", section_header_format)
            current_row += 1
            
            # Headers for normal payments with improved styling
            headers = ['SIRA NO', 'MÜŞTERİ ADI SOYADI', 'PROJE']
            week_day_names = week_data.get('week_day_names', [])
            headers.extend(week_day_names)
            headers.append('GENEL TOPLAM')
            
            for col, header in enumerate(headers):
                worksheet.write(current_row, col, header, table_header_format)
            current_row += 1
            
            # Normal payment data with improved formatting
            pivot_table = week_data.get('pivot', pd.DataFrame())
            if not pivot_table.empty:
                for idx, (customer_project, row) in enumerate(pivot_table.iterrows(), 1):
                    customer, project = customer_project
                    
                    # Row data with appropriate formatting
                    worksheet.write(current_row, 0, idx, data_cell_format)
                    worksheet.write(current_row, 1, customer, customer_cell_format)  # Left-aligned for readability
                    worksheet.write(current_row, 2, project, data_cell_format)
                    
                    # Daily amounts with currency formatting
                    col_idx = 3
                    row_total = 0
                    for date_col in pivot_table.columns:
                        amount = row[date_col] if pd.notna(row[date_col]) else 0
                        if amount > 0:
                            worksheet.write(current_row, col_idx, amount, currency_format)
                        else:
                            worksheet.write(current_row, col_idx, "", data_cell_format)
                        row_total += amount
                        col_idx += 1
                    
                    # Row total with highlighting
                    worksheet.write(current_row, col_idx, row_total, total_row_format)
                    current_row += 1
            
            current_row += 2
            
            # Check Payments Section (if exists) with improved formatting
            if check_data and not check_data.get('pivot', pd.DataFrame()).empty:
                worksheet.merge_range(current_row, 0, current_row, 10, 
                                    "ÇEK TAHSİLATLARI", section_header_format)
                current_row += 1
                
                # Check headers with improved styling
                for col, header in enumerate(headers):
                    worksheet.write(current_row, col, header, table_header_format)
                current_row += 1
                
                # Check payment data with improved formatting
                check_pivot = check_data.get('pivot', pd.DataFrame())
                for idx, (customer_project, row) in enumerate(check_pivot.iterrows(), 1):
                    customer, project = customer_project
                    
                    worksheet.write(current_row, 0, idx, data_cell_format)
                    worksheet.write(current_row, 1, customer, customer_cell_format)  # Left-aligned for readability
                    worksheet.write(current_row, 2, project, data_cell_format)
                    
                    col_idx = 3
                    row_total = 0
                    for date_col in check_pivot.columns:
                        amount = row[date_col] if pd.notna(row[date_col]) else 0
                        if amount > 0:
                            worksheet.write(current_row, col_idx, amount, currency_format)
                        else:
                            worksheet.write(current_row, col_idx, "", data_cell_format)
                        row_total += amount
                        col_idx += 1
                    
                    worksheet.write(current_row, col_idx, row_total, total_row_format)
                    current_row += 1
            
            current_row += 2
            
            # Analysis Tables (same as main export but for this week only)
            self._add_single_week_analysis_to_excel(worksheet, current_row, payments, start_date, end_date, workbook)
            
            # Set optimized column widths matching UI layout
            worksheet.set_column(0, 0, 8)   # SIRA NO - narrow for serial numbers
            worksheet.set_column(1, 1, 25)  # MÜŞTERİ ADI SOYADI - wide for customer names (matches UI's 200px)
            worksheet.set_column(2, 2, 15)  # PROJE - medium width for project names
            
            # Daily columns - optimized width (matches UI's 100px equivalent)
            daily_cols = len(week_day_names)
            for i in range(daily_cols):
                worksheet.set_column(3 + i, 3 + i, 12)
            
            # Total column
            worksheet.set_column(3 + daily_cols, 3 + daily_cols, 15)  # GENEL TOPLAM
            
            # Set row heights for better readability
            worksheet.set_row(0, 25)  # Main title row
            worksheet.set_row(1, 20)  # Week title row
                
        except Exception as e:
            logger.error(f"Failed to export single week to Excel: {e}")
            raise
    
    def _add_single_week_analysis_to_excel(self, worksheet, start_row, payments, start_date, end_date, workbook):
        """Add analysis tables for a single week with improved formatting"""
        try:
            # Generate analysis data
            payment_type_analysis = self.generate_payment_type_analysis(payments, start_date, end_date)
            project_totals_analysis = self.generate_project_totals_analysis(payments, start_date, end_date)
            location_analysis = self.generate_location_analysis(payments, start_date, end_date)
            
            # Use consistent formatting with the main export
            section_header_format = workbook.add_format({
                'bold': True,
                'font_size': 14,
                'font_color': '#495057',
                'bg_color': '#F8F9FA',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'text_wrap': True
            })
            
            table_header_format = workbook.add_format({
                'bold': True,
                'font_weight': 600,
                'font_color': '#495057',
                'bg_color': '#F8F9FA',
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'text_wrap': True
            })
            
            data_cell_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'font_color': '#212529'
            })
            
            currency_format = workbook.add_format({
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'num_format': '₺#,##0.00',
                'font_color': '#212529'
            })
            
            current_row = start_row
            
            # Payment Type Analysis
            worksheet.merge_range(current_row, 0, current_row, 4, "ÖDEME TİPİ ANALİZİ", section_header_format)
            current_row += 1
            
            headers = ['Ödeme Nedeni', 'Toplam TL', 'Toplam USD']
            for col, header in enumerate(headers):
                worksheet.write(current_row, col, header, table_header_format)
            current_row += 1
            
            # Add payment type data
            for payment_type in ['Nakit', 'BANK_TRANSFER', 'Çek']:
                weekly_data = payment_type_analysis.get('weekly', {})
                tl_amount = sum(week_data.get(payment_type, {}).get('total_tl', 0) for week_data in weekly_data.values())
                usd_amount = sum(week_data.get(payment_type, {}).get('total_usd', 0) for week_data in weekly_data.values())
                
                worksheet.write(current_row, 0, payment_type, cell_format)
                worksheet.write(current_row, 1, f"₺{tl_amount:,.2f}", cell_format)
                worksheet.write(current_row, 2, usd_amount, currency_format)
                current_row += 1
            
            current_row += 2
            
            # Project Totals Analysis
            worksheet.merge_range(current_row, 0, current_row, 2, "PROJE TOPLAMLARI", section_header_format)
            current_row += 1
            
            headers = ['Proje', 'Toplam USD']
            for col, header in enumerate(headers):
                worksheet.write(current_row, col, header, table_header_format)
            current_row += 1
            
            # Add project data
            weekly_projects = project_totals_analysis.get('weekly', {})
            total_PROJECT_A = sum(week_data.get('PROJECT_A', 0) for week_data in weekly_projects.values() if isinstance(week_data.get('PROJECT_A'), (int, float)))
            total_PROJECT_B = sum(week_data.get('PROJECT_B', 0) for week_data in weekly_projects.values() if isinstance(week_data.get('PROJECT_B'), (int, float)))
            
            worksheet.write(current_row, 0, 'PROJECT_A', cell_format)
            worksheet.write(current_row, 1, total_PROJECT_A, currency_format)
            current_row += 1
            
            worksheet.write(current_row, 0, 'PROJECT_B', cell_format)
            worksheet.write(current_row, 1, total_PROJECT_B, currency_format)
            current_row += 1
            
            worksheet.write(current_row, 0, 'TOPLAM', header_format)
            worksheet.write(current_row, 1, total_PROJECT_A + total_PROJECT_B, currency_format)
            current_row += 2
            
            # Location Analysis
            worksheet.merge_range(current_row, 0, current_row, 4, "LOKASYON ANALİZİ", header_format)
            current_row += 1
            
            headers = ['Lokasyon', 'PROJECT_A USD', 'PROJECT_B USD', 'Toplam USD']
            for col, header in enumerate(headers):
                worksheet.write(current_row, col, header, header_format)
            current_row += 1
            
            # Add location data
            weekly_locations = location_analysis.get('weekly', {})
            locations = ['ÇARŞI', 'LOCATION_B', 'OFİS', 'BANKA HAVALESİ', 'A KASA ÇEK', 'B KASA ÇEK']
            
            for location in locations:
                location_data = {}
                for week_data in weekly_locations.values():
                    if location in week_data:
                        loc_data = week_data[location]
                        if isinstance(loc_data, dict):
                            location_data['PROJECT_A'] = location_data.get('PROJECT_A', 0) + loc_data.get('PROJECT_A', 0)
                            location_data['PROJECT_B'] = location_data.get('PROJECT_B', 0) + loc_data.get('PROJECT_B', 0)
                
                PROJECT_A_amount = location_data.get('PROJECT_A', 0)
                PROJECT_B_amount = location_data.get('PROJECT_B', 0)
                total_amount = PROJECT_A_amount + PROJECT_B_amount
                
                worksheet.write(current_row, 0, location, cell_format)
                worksheet.write(current_row, 1, PROJECT_A_amount, currency_format)
                worksheet.write(current_row, 2, PROJECT_B_amount, currency_format)
                worksheet.write(current_row, 3, total_amount, currency_format)
                current_row += 1
                
        except Exception as e:
            logger.error(f"Failed to add single week analysis to Excel: {e}")
            raise
    
    def export_to_pdf(self, payments: List[PaymentData], 
                     start_date: datetime, end_date: datetime, 
                     output_path: str, orientation: str = "landscape") -> None:
        """Export reports to PDF with Turkish character support and orientation options"""
        try:
            # Choose page size based on orientation
            if orientation.lower() == "landscape":
                pagesize = landscape(A4)
            else:
                pagesize = A4
            
            doc = SimpleDocTemplate(output_path, pagesize=pagesize)
            story = []
            
            # Main Title with proper Turkish character encoding
            title_text = "MODEL KUYUM MERKEZİ - MODEL SANAYİ MERKEZİ TAHSİLATLAR TABLOSU"
            title = Paragraph(title_text, self.title_style)
            story.append(title)
            
            # Date range and payment type subtitle
            date_range_str = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
            subtitle_text = f"{date_range_str} | BANK_TRANSFER - Nakit"
            subtitle = Paragraph(subtitle_text, self.subtitle_style)
            story.append(subtitle)
            story.append(Spacer(1, 20))
            
            # Generate reports
            customer_date_table = self.generate_customer_date_table(payments, start_date, end_date)
            daily_breakdown = self.generate_daily_usd_breakdown(payments, start_date, end_date)
            monthly_summary = self.generate_monthly_summary(payments)
            payment_type_summary = self.generate_payment_type_summary(payments)
            
            # Customer-Date tables (weekly separated)
            if customer_date_table:
                sorted_weeks = sorted(customer_date_table.keys())
                
                for week_idx, week_start in enumerate(sorted_weeks):
                    week_data = customer_date_table[week_start]
                    pivot = week_data['pivot']
                    tl_converted = week_data['tl_converted']
                    day_names = week_data['day_names']
                    week_range = week_data['week_range']
                    week_dates = week_data['week_dates']
                    week_day_names = week_data['week_day_names']
                    
                    if pivot.empty:
                        continue
                    
                    # Add week header with page break if needed
                    if week_idx > 0:
                        story.append(PageBreak())
                        story.append(Spacer(1, 20))
                    
                    week_header = Paragraph(f"<b>HAFTA {week_idx + 1}: {week_range}</b>", self.subtitle_style)
                    story.append(week_header)
                    story.append(Spacer(1, 10))
                    
                    # Create table with day names and dates
                    table_data = []
                    
                    # Day names row
                    day_row = ['', ''] + week_day_names + ['']
                    table_data.append(day_row)
                    
                    # Date headers row
                    header_row = ['Müşteri Adı Soyadı', 'Proje'] + week_dates + ['Genel Toplam']
                    table_data.append(header_row)
                    
                    # Data rows
                    for (customer, project), row_data in pivot.iterrows():
                        data_row = [customer, project]
                        for date_col in week_dates:
                            amount = row_data[date_col]
                            is_tl = tl_converted.loc[(customer, project), date_col] if date_col in tl_converted.columns else False
                            
                            if amount > 0:
                                # Mark TL converted amounts with special formatting
                                if is_tl:
                                    data_row.append(f"${amount:,.0f}*")  # Add asterisk for TL converted
                                else:
                                    data_row.append(f"${amount:,.0f}")
                            else:
                                data_row.append("")
                        
                        # Total column
                        total_amount = row_data['Genel Toplam']
                        if total_amount > 0:
                            data_row.append(f"${total_amount:,.0f}")
                        else:
                            data_row.append("")
                        
                        table_data.append(data_row)
                    
                    # Week summary row
                    summary_row = ['Hafta Toplamı', '']
                    for date_col in week_dates:
                        total = pivot[date_col].sum()
                        if total > 0:
                            summary_row.append(f"${total:,.0f}")
                        else:
                            summary_row.append("")
                    
                    week_total = pivot['Genel Toplam'].sum()
                    if week_total > 0:
                        summary_row.append(f"${week_total:,.0f}")
                    else:
                        summary_row.append("")
                    
                    table_data.append(summary_row)
                    
                    # Create and style the table with page break prevention
                    table = Table(table_data, repeatRows=2)  # Repeat both day names and headers
                    
                    # Calculate number of columns
                    num_cols = len(table_data[0])
                    
                    # Set table properties to prevent page breaks within the table
                    table.hAlign = 'LEFT'
                    table.vAlign = 'TOP'
                    
                    table.setStyle(TableStyle([
                        # Day names row styling
                        ('BACKGROUND', (0, 0), (num_cols-1, 0), colors.Color(0.85, 0.85, 0.85)),  # Light gray
                        ('FONTNAME', (0, 0), (num_cols-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (num_cols-1, 0), 8),
                        ('ALIGN', (0, 0), (num_cols-1, 0), 'CENTER'),
                        
                        # Header row styling
                        ('BACKGROUND', (0, 1), (num_cols-1, 1), colors.Color(0.31, 0.51, 0.74)),  # Blue header
                        ('TEXTCOLOR', (0, 1), (num_cols-1, 1), colors.whitesmoke),
                        ('FONTNAME', (0, 1), (num_cols-1, 1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 1), (num_cols-1, 1), 9),
                        ('ALIGN', (0, 1), (num_cols-1, 1), 'CENTER'),
                        ('VALIGN', (0, 1), (num_cols-1, 1), 'MIDDLE'),
                        ('BOTTOMPADDING', (0, 1), (num_cols-1, 1), 8),
                        
                        # Data rows styling
                        ('BACKGROUND', (0, 2), (num_cols-1, -2), colors.Color(0.95, 0.95, 0.95)),  # Light gray
                        ('FONTNAME', (0, 2), (num_cols-1, -2), 'Helvetica'),
                        ('FONTSIZE', (0, 2), (num_cols-1, -2), 8),
                        ('ALIGN', (0, 2), (1, -2), 'LEFT'),  # Customer and project columns left aligned
                        ('ALIGN', (2, 2), (num_cols-1, -2), 'RIGHT'),  # Amount columns right aligned
                        
                        # Summary row styling
                        ('BACKGROUND', (0, -1), (num_cols-1, -1), colors.Color(0.69, 0.77, 0.87)),  # Light blue
                        ('FONTNAME', (0, -1), (num_cols-1, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, -1), (num_cols-1, -1), 9),
                        ('ALIGN', (0, -1), (1, -1), 'LEFT'),
                        ('ALIGN', (2, -1), (num_cols-1, -1), 'RIGHT'),
                        
                        # Grid and borders
                        ('GRID', (0, 0), (num_cols-1, -1), 1, colors.black),
                        ('LINEBELOW', (0, 1), (num_cols-1, 1), 2, colors.black),  # Thick line under header
                        ('LINEABOVE', (0, -1), (num_cols-1, -1), 2, colors.black),  # Thick line above summary
                    ]))
                    
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                # Add legend for TL converted amounts
                legend = Paragraph("<i>* Amounts marked with asterisk (*) are converted from TL to USD</i>", 
                                 self.styles['Normal'])
                story.append(legend)
                story.append(Spacer(1, 20))
            
            # Additional reports sections
            if not daily_breakdown.empty:
                story.append(Paragraph("Günlük USD Dağılımı", self.styles['Heading2']))
                story.append(Spacer(1, 12))
                
                # Convert to table format
                table_data = [['Müşteri', 'Proje'] + list(daily_breakdown.columns)]
                for (customer, project), row in daily_breakdown.iterrows():
                    table_data.append([customer, project] + [f"${val:,.2f}" for val in row.values])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
                story.append(Spacer(1, 20))
            
            # Monthly summary table
            if not monthly_summary.empty:
                story.append(Paragraph("Aylık Kanal Dağılımı", self.styles['Heading2']))
                story.append(Spacer(1, 12))
                
                table_data = [['Kanal'] + list(monthly_summary.columns)]
                for channel, row in monthly_summary.iterrows():
                    table_data.append([channel] + [f"${val:,.2f}" for val in row.values])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(table)
            
            doc.build(story)
            logger.info(f"PDF report exported to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export PDF report: {e}")
            raise
    
    def export_to_word(self, payments: List[PaymentData], 
                      start_date: datetime, end_date: datetime, 
                      output_path: str) -> None:
        """Export reports to Word document"""
        try:
            doc = Document()
            
            # Main Title
            title_text = "MODEL KUYUM MERKEZİ - MODEL SANAYİ MERKEZİ TAHSİLATLAR TABLOSU"
            title = doc.add_heading(title_text, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date range and payment type subtitle
            date_range_str = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
            subtitle_text = f"{date_range_str} | BANK_TRANSFER - Nakit"
            subtitle = doc.add_heading(subtitle_text, level=1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Generate reports
            customer_date_table = self.generate_customer_date_table(payments, start_date, end_date)
            daily_breakdown = self.generate_daily_usd_breakdown(payments, start_date, end_date)
            monthly_summary = self.generate_monthly_summary(payments)
            payment_type_summary = self.generate_payment_type_summary(payments)
            
            # Customer-Date tables (weekly separated)
            if customer_date_table:
                sorted_weeks = sorted(customer_date_table.keys())
                
                for week_idx, week_start in enumerate(sorted_weeks):
                    week_data = customer_date_table[week_start]
                    pivot = week_data['pivot']
                    tl_converted = week_data['tl_converted']
                    day_names = week_data['day_names']
                    week_range = week_data['week_range']
                    week_dates = week_data['week_dates']
                    week_day_names = week_data['week_day_names']
                    
                    if pivot.empty:
                        continue
                    
                    # Add week header
                    if week_idx > 0:
                        doc.add_paragraph("")  # Add spacing
                    
                    week_header = doc.add_heading(f'HAFTA {week_idx + 1}: {week_range}', level=2)
                    week_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Create table
                    num_cols = len(week_dates) + 3  # Customer, Project, dates, Total
                    table = doc.add_table(rows=3, cols=num_cols)  # Day names, headers, and first data row
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Day names row
                    day_cells = table.rows[0].cells
                    day_cells[0].text = ''
                    day_cells[1].text = ''
                    for i, date_col in enumerate(week_dates):
                        day_name = day_names.get(date_col, '')
                        day_cells[i + 2].text = day_name
                    day_cells[-1].text = ''
                    
                    # Make day names row bold
                    for cell in day_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Header row
                    hdr_cells = table.rows[1].cells
                    hdr_cells[0].text = 'Müşteri Adı Soyadı'
                    hdr_cells[1].text = 'Proje'
                    for i, date_col in enumerate(week_dates):
                        hdr_cells[i + 2].text = date_col
                    hdr_cells[-1].text = 'Genel Toplam'
                    
                    # Make header row bold
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Remove the pre-added third row (we'll add data rows properly)
                    table._tbl.remove(table.rows[2]._tr)
                    
                    # Data rows
                    for (customer, project), row_data in pivot.iterrows():
                        row_cells = table.add_row().cells
                        row_cells[0].text = customer
                        row_cells[1].text = project
                        
                        for i, date_col in enumerate(week_dates):
                            amount = row_data[date_col]
                            is_tl = tl_converted.loc[(customer, project), date_col] if date_col in tl_converted.columns else False
                            
                            if amount > 0:
                                # Mark TL converted amounts
                                if is_tl:
                                    row_cells[i + 2].text = f"${amount:,.0f}*"
                                else:
                                    row_cells[i + 2].text = f"${amount:,.0f}"
                            else:
                                row_cells[i + 2].text = ""
                        
                        # Total column
                        total_amount = row_data['Genel Toplam']
                        if total_amount > 0:
                            row_cells[-1].text = f"${total_amount:,.0f}"
                        else:
                            row_cells[-1].text = ""
                    
                    # Add week summary row
                    summary_cells = table.add_row().cells
                    summary_cells[0].text = 'Hafta Toplamı'
                    summary_cells[1].text = ''
                    
                    for i, date_col in enumerate(week_dates):
                        total = pivot[date_col].sum()
                        if total > 0:
                            summary_cells[i + 2].text = f"${total:,.0f}"
                        else:
                            summary_cells[i + 2].text = ""
                    
                    # Week total
                    week_total = pivot['Genel Toplam'].sum()
                    if week_total > 0:
                        summary_cells[-1].text = f"${week_total:,.0f}"
                    else:
                        summary_cells[-1].text = ""
                    
                    # Make summary row bold
                    for cell in summary_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    doc.add_paragraph("")  # Add spacing after table
                
                # Add legend for TL converted amounts
                legend = doc.add_paragraph()
                legend_run = legend.add_run("* Amounts marked with asterisk (*) are converted from TL to USD")
                legend_run.italic = True
                doc.add_paragraph("")
            
            # Additional reports sections
            if not daily_breakdown.empty:
                doc.add_heading('Günlük USD Dağılımı', level=1)
                
                table = doc.add_table(rows=1, cols=len(daily_breakdown.columns) + 2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Müşteri'
                hdr_cells[1].text = 'Proje'
                for i, col in enumerate(daily_breakdown.columns):
                    hdr_cells[i + 2].text = str(col)
                
                # Data rows
                for (customer, project), row in daily_breakdown.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = customer
                    row_cells[1].text = project
                    for i, val in enumerate(row.values):
                        row_cells[i + 2].text = f"${val:,.2f}"
                
                doc.add_paragraph("")
            
            # Monthly summary table
            if not monthly_summary.empty:
                doc.add_heading('Aylık Kanal Dağılımı', level=1)
                
                table = doc.add_table(rows=1, cols=len(monthly_summary.columns) + 1)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Kanal'
                for i, col in enumerate(monthly_summary.columns):
                    hdr_cells[i + 1].text = str(col)
                
                # Data rows
                for channel, row in monthly_summary.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = channel
                    for i, val in enumerate(row.values):
                        row_cells[i + 1].text = f"${val:,.2f}"
            
            doc.save(output_path)
            logger.info(f"Word report exported to {output_path}")
        except Exception as e:
            logger.error(f"Failed to export Word report: {e}")
            raise
    
    def generate_payment_type_analysis(self, payments: List[PaymentData], 
                                     start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Generate payment type analysis (BANK_TRANSFER, Nakit, Çek) with TL/USD totals"""
        # Filter payments by date range using payment date consistently for all payment types
        filtered_payments = []
        for p in payments:
            if p.date:
                # Convert dates to date objects for comparison
                start_date_obj = start_date.date() if hasattr(start_date, 'date') else start_date
                end_date_obj = end_date.date() if hasattr(end_date, 'date') else end_date
                
                # Use payment date for filtering all payment types consistently
                payment_date = p.date.date() if hasattr(p.date, 'date') else p.date
                if start_date_obj <= payment_date <= end_date_obj:
                    filtered_payments.append(p)
        
        if not filtered_payments:
            return {'weekly': {}, 'monthly': {}}
        
        # Process payments by type
        payment_types = {
            'BANK_TRANSFER': [],
            'Nakit': [],
            'Çek': []
        }
        
        for payment in filtered_payments:
            # Use the payment_type field we already calculated in PaymentData
            payment_type = getattr(payment, 'payment_type', 'Diğer')
            
            # Ensure we have the correct categories
            if payment_type not in ['Nakit', 'BANK_TRANSFER', 'Çek']:
                # Fallback logic if payment_type is not set correctly
                tahsilat_sekli = payment.tahsilat_sekli.upper() if payment.tahsilat_sekli else ''
                account_name = payment.account_name.upper() if payment.account_name else ''
                
                if payment.is_check_payment or 'ÇEK' in tahsilat_sekli:
                    payment_type = 'Çek'
                elif 'YAPI KREDİ' in account_name or 'HAVALE' in tahsilat_sekli:
                    payment_type = 'BANK_TRANSFER'
                elif 'KASA' in account_name or 'NAKİT' in tahsilat_sekli:
                    payment_type = 'Nakit'
                else:
                    payment_type = 'BANK_TRANSFER'  # Default
            
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            payment_types[payment_type].append({
                'date': payment.date,
                'amount_tl': payment.amount if payment.is_tl_payment else 0,
                'amount_usd': usd_amount,
                'week_start': payment.date - timedelta(days=payment.date.weekday()),
                'month': payment.date.strftime('%Y-%m')
            })
        
        # Generate weekly analysis
        weekly_analysis = {}
        for week_start in pd.date_range(start=start_date, end=end_date, freq='W-MON'):
            week_end = week_start + timedelta(days=6)
            if week_end > end_date:
                week_end = end_date
            
            week_data = {
                'BANK_TRANSFER': {'tl_total': 0, 'usd_total': 0},
                'Nakit': {'tl_total': 0, 'usd_total': 0},
                'Çek': {'tl_total': 0, 'usd_total': 0}
            }
            
            for payment_type, type_payments in payment_types.items():
                for payment in type_payments:
                    if week_start <= payment['date'] <= week_end:
                        week_data[payment_type]['tl_total'] += payment['amount_tl']
                        week_data[payment_type]['usd_total'] += payment['amount_usd']
            
            # Calculate totals
            week_data['Genel Toplam'] = {
                'tl_total': sum(data['tl_total'] for data in week_data.values() if isinstance(data, dict)),
                'usd_total': sum(data['usd_total'] for data in week_data.values() if isinstance(data, dict))
            }
            
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_analysis[week_key] = week_data
        
        # Generate monthly analysis
        monthly_analysis = {
            'BANK_TRANSFER': {'tl_total': 0, 'usd_total': 0},
            'Nakit': {'tl_total': 0, 'usd_total': 0},
            'Çek': {'tl_total': 0, 'usd_total': 0}
        }
        
        for payment_type, type_payments in payment_types.items():
            for payment in type_payments:
                monthly_analysis[payment_type]['tl_total'] += payment['amount_tl']
                monthly_analysis[payment_type]['usd_total'] += payment['amount_usd']
        
        monthly_analysis['Genel Toplam'] = {
            'tl_total': sum(data['tl_total'] for data in monthly_analysis.values() if isinstance(data, dict)),
            'usd_total': sum(data['usd_total'] for data in monthly_analysis.values() if isinstance(data, dict))
        }
        
        return {
            'weekly': weekly_analysis,
            'monthly': monthly_analysis
        }
    
    def generate_project_totals_analysis(self, payments: List[PaymentData], 
                                       start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Generate PROJECT_A/PROJECT_B project totals analysis for weekly and monthly USD payments"""
        # Filter payments by date range (more flexible date comparison)
        filtered_payments = []
        for p in payments:
            if p.date:
                # Convert dates to date objects for comparison
                payment_date = p.date.date() if hasattr(p.date, 'date') else p.date
                start_date_obj = start_date.date() if hasattr(start_date, 'date') else start_date
                end_date_obj = end_date.date() if hasattr(end_date, 'date') else end_date
                
                if start_date_obj <= payment_date <= end_date_obj:
                    filtered_payments.append(p)
        
        if not filtered_payments:
            return {'weekly': {}, 'monthly': {}}
        
        # Process payments by project
        project_payments = {
            'PROJECT_A': [],
            'PROJECT_B': []
        }
        
        for payment in filtered_payments:
            project_name = payment.project_name.upper() if payment.project_name else ''
            
            # Determine project type based on actual project names
            if 'PROJECT_A' in project_name:
                project_type = 'PROJECT_A'
            elif 'PROJECT_B' in project_name:
                project_type = 'PROJECT_B'
            elif 'KUYUM' in project_name or 'KIYIM' in project_name:
                project_type = 'PROJECT_A'  # COMPANY_A
            elif 'SANAYİ' in project_name or 'SANAYI' in project_name:
                project_type = 'PROJECT_B'  # COMPANY_B
            else:
                # Default classification
                project_type = 'PROJECT_A'  # Default to PROJECT_A
            
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            project_payments[project_type].append({
                'date': payment.date,
                'amount_usd': usd_amount,
                'week_start': payment.date - timedelta(days=payment.date.weekday()),
                'month': payment.date.strftime('%Y-%m')
            })
        
        # Generate weekly analysis
        weekly_analysis = {}
        for week_start in pd.date_range(start=start_date, end=end_date, freq='W-MON'):
            week_end = week_start + timedelta(days=6)
            if week_end > end_date:
                week_end = end_date
            
            week_data = {
                'PROJECT_A': 0,
                'PROJECT_B': 0
            }
            
            for project_type, project_payment_list in project_payments.items():
                for payment in project_payment_list:
                    if week_start <= payment['date'] <= week_end:
                        week_data[project_type] += payment['amount_usd']
            
            week_data['TOPLAM'] = week_data['PROJECT_A'] + week_data['PROJECT_B']
            
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_analysis[week_key] = week_data
        
        # Generate monthly analysis
        monthly_analysis = {
            'PROJECT_A': 0,
            'PROJECT_B': 0
        }
        
        for project_type, project_payment_list in project_payments.items():
            for payment in project_payment_list:
                monthly_analysis[project_type] += payment['amount_usd']
        
        monthly_analysis['TOPLAM'] = monthly_analysis['PROJECT_A'] + monthly_analysis['PROJECT_B']
        
        return {
            'weekly': weekly_analysis,
            'monthly': monthly_analysis
        }
    
    def generate_location_analysis(self, payments: List[PaymentData], 
                                 start_date: datetime, end_date: datetime) -> Dict[str, Any]:
        """Generate location-based payment analysis (Çarşı, LOCATION_B, LOCATION_C, BANK_TRANSFER, A Kasa Çek, B Kasa Çek)"""
        # Filter payments by date range (more flexible date comparison)
        filtered_payments = []
        for p in payments:
            if p.date:
                # Convert dates to date objects for comparison
                payment_date = p.date.date() if hasattr(p.date, 'date') else p.date
                start_date_obj = start_date.date() if hasattr(start_date, 'date') else start_date
                end_date_obj = end_date.date() if hasattr(end_date, 'date') else end_date
                
                if start_date_obj <= payment_date <= end_date_obj:
                    filtered_payments.append(p)
        
        if not filtered_payments:
            return {'weekly': {}, 'monthly': {}}
        
        # Process payments by location and check type
        location_payments = {
            'ÇARŞI': {'PROJECT_A': [], 'PROJECT_B': []},
            'LOCATION_B': {'PROJECT_A': [], 'PROJECT_B': []},
            'OFİS': {'PROJECT_A': [], 'PROJECT_B': []},
            'BANKA HAVALESİ': {'PROJECT_A': [], 'PROJECT_B': []},
            'A KASA ÇEK': {'PROJECT_A': [], 'PROJECT_B': []},
            'B KASA ÇEK': {'PROJECT_A': [], 'PROJECT_B': []}
        }
        
        for payment in filtered_payments:
            # Determine location and check type from Hesap Adı
            account_name = payment.account_name.upper() if payment.account_name else ''
            
            # Check for location keywords (handle encoding issues)
            if 'LOCATION_B' in account_name or 'KUYUMCU' in account_name:
                location = 'LOCATION_B'
            elif 'ÇARŞI' in account_name or 'CARSI' in account_name or 'ÇARÞI' in account_name or 'CARÞI' in account_name:
                location = 'ÇARŞI'
            elif ('OFİS' in account_name or 'LOCATION_C' in account_name or 'MERKEZ' in account_name or 
                  'OFÝS' in account_name or 'OFİS' in account_name):
                location = 'OFİS'
            elif ('HAVALE' in account_name or 'TRANSFER' in account_name or 'BANKA' in account_name or 
                  'KREDİ' in account_name or 'KREDÝ' in account_name or 'KREDI' in account_name):
                location = 'BANKA HAVALESİ'
            elif 'ÇEK' in account_name or 'CEK' in account_name or 'CHECK' in account_name or payment.is_check_payment:
                # Determine check type (A or B)
                if 'A' in account_name or 'RESMİ' in account_name or 'RESMI' in account_name:
                    location = 'A KASA ÇEK'
                elif 'B' in account_name or 'GAYRİRESMİ' in account_name or 'GAYRIRESMI' in account_name:
                    location = 'B KASA ÇEK'
                else:
                    # Default to A type for checks
                    location = 'A KASA ÇEK'
            else:
                # Default classification
                if payment.is_check_payment:
                    location = 'A KASA ÇEK'
                else:
                    location = 'BANKA HAVALESİ'
            
            # Determine project type
            project_name = payment.project_name.upper() if payment.project_name else ''
            if 'PROJECT_A' in project_name:
                project_type = 'PROJECT_A'
            elif 'PROJECT_B' in project_name:
                project_type = 'PROJECT_B'
            elif 'KUYUM' in project_name or 'KIYIM' in project_name:
                project_type = 'PROJECT_A'  # COMPANY_A
            elif 'SANAYİ' in project_name or 'SANAYI' in project_name:
                project_type = 'PROJECT_B'  # COMPANY_B
            else:
                project_type = 'PROJECT_A'  # Default to PROJECT_A
            
            # Use the already converted USD amount from PaymentData
            usd_amount = payment.usd_amount if payment.usd_amount > 0 else payment.amount
            
            location_payments[location][project_type].append({
                'date': payment.date,
                'amount_usd': usd_amount,
                'week_start': payment.date - timedelta(days=payment.date.weekday()),
                'month': payment.date.strftime('%Y-%m')
            })
        
        # Generate weekly analysis
        weekly_analysis = {}
        for week_start in pd.date_range(start=start_date, end=end_date, freq='W-MON'):
            week_end = week_start + timedelta(days=6)
            if week_end > end_date:
                week_end = end_date
            
            week_data = {
                'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
                'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0}
            }
            
            for location, project_data in location_payments.items():
                for project_type, payment_list in project_data.items():
                    for payment in payment_list:
                        if week_start <= payment['date'] <= week_end:
                            week_data[location][project_type] += payment['amount_usd']
            
            # Calculate totals
            week_data['TOPLAM'] = {
                'PROJECT_A': sum(data['PROJECT_A'] for data in week_data.values() if isinstance(data, dict)),
                'PROJECT_B': sum(data['PROJECT_B'] for data in week_data.values() if isinstance(data, dict))
            }
            week_data['TOPLAM']['GENEL'] = week_data['TOPLAM']['PROJECT_A'] + week_data['TOPLAM']['PROJECT_B']
            
            week_key = week_start.strftime('%Y-%m-%d')
            weekly_analysis[week_key] = week_data
        
        # Generate monthly analysis
        monthly_analysis = {
            'ÇARŞI': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'LOCATION_B': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'OFİS': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'BANKA HAVALESİ': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'A KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0},
            'B KASA ÇEK': {'PROJECT_A': 0, 'PROJECT_B': 0}
        }
        
        for location, project_data in location_payments.items():
            for project_type, payment_list in project_data.items():
                for payment in payment_list:
                    monthly_analysis[location][project_type] += payment['amount_usd']
        
        # Calculate monthly totals
        monthly_analysis['TOPLAM'] = {
            'PROJECT_A': sum(data['PROJECT_A'] for data in monthly_analysis.values() if isinstance(data, dict)),
            'PROJECT_B': sum(data['PROJECT_B'] for data in monthly_analysis.values() if isinstance(data, dict))
        }
        monthly_analysis['TOPLAM']['GENEL'] = monthly_analysis['TOPLAM']['PROJECT_A'] + monthly_analysis['TOPLAM']['PROJECT_B']
        
        return {
            'weekly': weekly_analysis,
            'monthly': monthly_analysis
        }

# Convenience functions
def generate_all_reports(payments: List[PaymentData], 
                        start_date: datetime, end_date: datetime,
                        output_dir: str = "reports") -> Dict[str, str]:
    """Generate all reports in all formats"""
    output_path = Path(output_dir)
    output_path.mkdir(exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    generator = ReportGenerator()
    
    # Generate reports
    excel_path = output_path / f"tahsilat_raporu_{timestamp}.xlsx"
    pdf_path = output_path / f"tahsilat_raporu_{timestamp}.pdf"
    word_path = output_path / f"tahsilat_raporu_{timestamp}.docx"
    
    generator.export_to_excel(payments, start_date, end_date, str(excel_path))
    generator.export_to_pdf(payments, start_date, end_date, str(pdf_path))
    generator.export_to_word(payments, start_date, end_date, str(word_path))
    
    return {
        'excel': str(excel_path),
        'pdf': str(pdf_path),
        'word': str(word_path)
    }
